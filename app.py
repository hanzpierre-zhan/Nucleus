import os
import glob
import json
import time
import logging
import tempfile
import mimetypes
import pandas as pd
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, send_from_directory, Response
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from functools import wraps
import io
from datetime import datetime, timedelta
from collections import Counter
from PIL import Image, ImageOps

# Setup Flask application
app = Flask(__name__)
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DATABASE_URL = os.environ.get('DATABASE_URL')
if DATABASE_URL:
    app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
else:
    db_path = os.path.join(BASE_DIR, "nucleus.db")
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {'pool_pre_ping': True}
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'nucleus_dev_key_change_me')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024 # 50 MB
app.config['EVIDENCIA_DIR'] = os.path.join(BASE_DIR, 'static', 'evidencia')
os.makedirs(app.config['EVIDENCIA_DIR'], exist_ok=True)
app.config['EVIDENCIA_MAX_LADO'] = int(os.environ.get('EVIDENCIA_MAX_LADO', 1280))
app.config['EVIDENCIA_CALIDAD'] = int(os.environ.get('EVIDENCIA_CALIDAD', 80))
# Backblaze B2 (opcional): si se configuran estas variables, las fotos se suben a B2.
app.config['B2_ENDPOINT_URL'] = os.environ.get('B2_ENDPOINT_URL', '')
app.config['B2_KEY_ID'] = os.environ.get('B2_KEY_ID', '')
app.config['B2_APP_KEY'] = os.environ.get('B2_APP_KEY', '')
app.config['B2_BUCKET'] = os.environ.get('B2_BUCKET', '')
app.config['B2_REGION'] = os.environ.get('B2_REGION', 'us-west-004')

db = SQLAlchemy(app)

# --- MODELS ---
class Usuario(db.Model):
    __tablename__ = 'usuarios'
    id = db.Column(db.Integer, primary_key=True)
    nombre = db.Column(db.String(100), default='')
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    rol = db.Column(db.String(20), default='supervisor') # 'admin', 'supervisor', 'gestor'

class Proyecto(db.Model):
    __tablename__ = 'proyectos'
    id = db.Column(db.Integer, primary_key=True)
    nombre = db.Column(db.String(100), unique=True, nullable=False)
    descripcion = db.Column(db.String(200))
    icono = db.Column(db.String(50), default='fa-folder-open', nullable=False)

class AppConfig(db.Model):
    __tablename__ = 'app_config'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    clave = db.Column(db.String(50), nullable=False)
    valor = db.Column(db.Text, nullable=False)
    __table_args__ = (db.UniqueConstraint('proyecto_id', 'clave', name='_proj_clave_uc'),)

class NucleusData(db.Model):
    __tablename__ = 'nucleus_data'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    key_value = db.Column(db.String(100), nullable=False, index=True)
    data_json = db.Column(db.Text, nullable=False)
    __table_args__ = (db.UniqueConstraint('proyecto_id', 'key_value', name='_proj_key_uc'),)

class NucleusHistory(db.Model):
    __tablename__ = 'nucleus_history'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    key_value = db.Column(db.String(100), nullable=False, index=True)
    data_json = db.Column(db.Text, nullable=False)
    fecha_consolidado = db.Column(db.DateTime, default=datetime.utcnow)

class FiltroMaestro(db.Model):
    __tablename__ = 'filtros_maestros'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    columna = db.Column(db.String(100), nullable=False)
    valor = db.Column(db.String(100), nullable=False)
    __table_args__ = (db.UniqueConstraint('proyecto_id', 'columna', 'valor', name='_proj_filtro_uc'),)

class TablaMaestra(db.Model):
    __tablename__ = 'tablas_maestras'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    columna_criterio = db.Column(db.String(100), nullable=False)
    valor_criterio = db.Column(db.String(100), nullable=False)
    nueva_columna = db.Column(db.String(100), nullable=False)
    nuevo_valor = db.Column(db.String(100), nullable=False)

class ReglaEstadoManual(db.Model):
    __tablename__ = 'reglas_estado_manual'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    columna_criterio = db.Column(db.String(100), nullable=False)
    valor_criterio = db.Column(db.String(100), nullable=False)
    columna_manual = db.Column(db.String(100), nullable=False)
    nuevo_valor = db.Column(db.String(100), nullable=False)

class AccesoProyecto(db.Model):
    __tablename__ = 'accesos_proyecto'
    id = db.Column(db.Integer, primary_key=True)
    usuario_id = db.Column(db.Integer, db.ForeignKey('usuarios.id'), nullable=False)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    # restricciones: JSON string ej: {"JEFATURA": ["LIMA", "CALLAO"]}
    restricciones = db.Column(db.Text, default='{}') 
    __table_args__ = (db.UniqueConstraint('usuario_id', 'proyecto_id', name='_user_proj_uc'),)

class KpiConfig(db.Model):
    __tablename__ = 'kpi_configs'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    nombre = db.Column(db.String(100), nullable=False)
    col_inicio = db.Column(db.String(100), nullable=False)
    restar_contra = db.Column(db.String(20), default='HOY') # 'HOY', 'COLUMNA'
    col_fin = db.Column(db.String(100), nullable=True)
    tipo = db.Column(db.String(20), default='DILACION')

class HistorialCambios(db.Model):
    __tablename__ = 'historial_cambios'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    usuario_id = db.Column(db.Integer, db.ForeignKey('usuarios.id'), nullable=True)
    username = db.Column(db.String(50), nullable=False)
    key_value = db.Column(db.String(100), nullable=False, index=True)
    campo_modificado = db.Column(db.String(100), nullable=False)
    valor_anterior = db.Column(db.Text, nullable=True)
    valor_nuevo = db.Column(db.Text, nullable=True)
    fecha = db.Column(db.DateTime, default=datetime.utcnow)

class Tecnico(db.Model):
    __tablename__ = 'tecnicos'
    id = db.Column(db.Integer, primary_key=True)
    proyecto_id = db.Column(db.Integer, db.ForeignKey('proyectos.id'), nullable=False)
    nombre = db.Column(db.String(120), nullable=False)
    contrata = db.Column(db.String(120), default='')
    especialidad = db.Column(db.String(120), default='')
    telefono = db.Column(db.String(30), default='')

# Auth Decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# --- HELPERS ---
@app.route('/api/admin/column_values')
@login_required
def api_column_values():
    pid_arg = request.args.get('pid')
    pid = pid_arg if pid_arg else session.get('current_proyecto_id')
    col = request.args.get('col')
    if not pid or not col: return jsonify([])
    
    # Get unique values from NucleusData
    rows = NucleusData.query.filter_by(proyecto_id=pid).all()
    vals = set()
    for r in rows:
        d = json.loads(r.data_json)
        v = d.get(col)
        if v: vals.add(str(v).strip())
    
    return jsonify(sorted(list(vals)))

def inject_kpis(pid, rows):
    configs = KpiConfig.query.filter_by(proyecto_id=pid).all()
    if not configs: return rows, {}
    
    hoy = datetime.now()
    kpi_meta = {}
    
    # Pre-calculate rules
    for kpi in configs:
        if kpi.tipo == 'ACUMULADO':
            col_raw = kpi.col_inicio
            # ignore empty strings from splitting
            cols = [c.strip() for c in col_raw.split(',') if c.strip()]
            
            if not cols: continue

            # Handle multiple filters stored in col_fin as JSON
            filters = []
            try:
                if kpi.col_fin and (kpi.col_fin.startswith('[') or kpi.col_fin.startswith('{')):
                    filters = json.loads(kpi.col_fin)
                    if not isinstance(filters, list): filters = []
                elif kpi.restar_contra and kpi.restar_contra != 'HOY':
                    # Legacy single filter support
                    filters = [{"col": kpi.restar_contra, "val": kpi.col_fin}]
            except:
                filters = []

            combined_vals = []
            for r in rows:
                # Apply ALL filters (AND logic)
                matches_all = True
                for f in filters:
                    f_col = f.get('col')
                    f_val = f.get('val')
                    if f_col:
                        # rows are list of dicts, use r.get()
                        if str(r.get(f_col, '')).strip() != str(f_val).strip():
                            matches_all = False
                            break
                
                if matches_all:
                    vals = [str(r.get(c, '')).strip() for c in cols]
                    if all(vals):
                        combined_vals.append(" | ".join(vals))
            
            if combined_vals:
                counts = Counter(combined_vals)
                
                # --- NEW TOP 4 RANKING LOGIC ---
                sorted_keys = sorted(counts.keys(), key=lambda x: counts[x], reverse=True)
                ranks = {}
                for i, k in enumerate(sorted_keys[:4]):
                    ranks[k] = i + 1  # Rank 1, 2, 3, 4
                
                # Use KPI ID to avoid collisions
                kpi_meta[kpi.id] = {
                    'counts': dict(counts),
                    'ranks': ranks,
                    'max': max(counts.values()) if counts else 0,
                    'cols_involved': cols,
                    'filters': filters
                }
        elif kpi.tipo == 'RESALTADO':
            # Highlight rules: { "COLUMN": {"VALUE": "STYLE"} }
            if 'resaltadores' not in kpi_meta: kpi_meta['resaltadores'] = {}
            col = kpi.col_inicio
            val = kpi.col_fin # ITEM to highlight stored here
            if col not in kpi_meta['resaltadores']: kpi_meta['resaltadores'][col] = {}
            kpi_meta['resaltadores'][col][val] = 'hit' # Mark for highlighting

    for kpi in configs:
        if kpi.tipo != 'DILACION': continue
        
        for row in rows:
            val_inicio = row.get(kpi.col_inicio)
            if not val_inicio:
                row[f"KPI_{kpi.nombre}"] = None
                continue
                
            # Try to parse date using pandas for robustness
            try:
                # Use pandas to parse almost any common date format
                f_inicio = pd.to_datetime(str(val_inicio).strip())
                if pd.isna(f_inicio): f_inicio = None
            except:
                f_inicio = None
            
            if not f_inicio:
                row[f"KPI_{kpi.nombre}"] = None
                continue
                
            target_date = pd.to_datetime(hoy)
            if kpi.restar_contra == 'COLUMNA' and kpi.col_fin:
                val_fin = row.get(kpi.col_fin)
                if val_fin:
                    try:
                        f_fin = pd.to_datetime(str(val_fin).strip())
                        if not pd.isna(f_fin):
                            target_date = f_fin
                    except: pass
            
            diff = target_date - f_inicio
            row[f"KPI_{kpi.nombre}"] = max(0, diff.days)
            
    return rows, kpi_meta

def apply_data_restrictions(data_list, res_obj):
    """
    Applies role-based data filtering based on the res_obj.
    Does case-insensitive comparison and ignores missing columns.
    """
    if not res_obj:
        return data_list
        
    filtered = []
    
    # Pre-process restrictions to be robust (uppercase lists)
    parsed_res = {}
    for col, vals in res_obj.items():
        if not vals: continue
        # Ensure it's a list and uppercase all elements
        if isinstance(vals, list):
            parsed_res[col] = [str(v).strip().upper() for v in vals]
        else:
            parsed_res[col] = [str(vals).strip().upper()]
    
    # If after parsing it's empty, no restrictions apply
    if not parsed_res:
        return data_list
        
    for d in data_list:
        keep = True
        for col_name, allowed_vals in parsed_res.items():
            if col_name not in d: continue # Flexible filtering
            
            val_in_row = str(d.get(col_name, '')).strip().upper()
            if val_in_row not in allowed_vals:
                keep = False; break
        if keep:
            filtered.append(d)
            
    return filtered


# --- DB INIT & MIGRATION ---
with app.app_context():
    is_sqlite = db.engine.dialect.name == 'sqlite'
    # Detect if we need to migrate or just create
    from sqlalchemy import inspect
    inspector = inspect(db.engine)
    tables = inspector.get_table_names()
    
    # Check if _old tables already exist from a failed prior run
    has_old = 'nucleus_data_old' in tables
    
    # If it's the old schema (no project_id in nucleus_data)
    needs_migration = False
    if 'nucleus_data' in tables:
        cols = [c['name'] for c in inspector.get_columns('nucleus_data')]
        if 'proyecto_id' not in cols:
            needs_migration = True

    if is_sqlite and needs_migration and not has_old:
        print("Migrating database to Multi-Project schema...")
        # Simple migration: Rename old and copy
        db.session.execute(db.text("ALTER TABLE nucleus_data RENAME TO nucleus_data_old"))
        db.session.execute(db.text("ALTER TABLE app_config RENAME TO app_config_old"))
        db.session.execute(db.text("ALTER TABLE filtros_maestros RENAME TO filtros_maestros_old"))
        db.session.execute(db.text("ALTER TABLE tablas_maestras RENAME TO tablas_maestras_old"))
        db.session.commit()
        has_old = True # Now we have them

    db.create_all()

    # Migration: Add 'nombre' column to usuarios if missing (must run before any Usuario query)
    try:
        ucols = [c['name'] for c in inspect(db.engine).get_columns('usuarios')]
        if 'nombre' not in ucols:
            db.session.execute(db.text("ALTER TABLE usuarios ADD COLUMN nombre VARCHAR(100) DEFAULT ''"))
            db.session.commit()
            print("Added 'nombre' column to usuarios")
    except Exception as e:
        print("Warning: could not add nombre column:", e)
    
    # Create Default Admin if none
    if not Usuario.query.first():
        admin = Usuario(username='admin', password_hash=generate_password_hash('admin123'), rol='admin')
        db.session.add(admin)
        db.session.commit()
        
    # Create Default Project "Pangeaco" ONLY when the DB is completely empty (fresh install)
    if not Proyecto.query.first():
        pangeaco = Proyecto(nombre='Pangeaco', descripcion='Proyecto inicial migrado')
        db.session.add(pangeaco)
        db.session.commit()
    else:
        # Reuse an existing project as target for any legacy migration
        pangeaco = Proyecto.query.first()
        
    if is_sqlite and has_old:
        pid = pangeaco.id
        print("Restoring data from old tables...")
        # Move data from old tables
        try:
            db.session.execute(db.text(f"INSERT OR IGNORE INTO nucleus_data (proyecto_id, key_value, data_json) SELECT {pid}, key_value, data_json FROM nucleus_data_old"))
            db.session.execute(db.text(f"INSERT OR IGNORE INTO app_config (proyecto_id, clave, valor) SELECT {pid}, clave, valor FROM app_config_old"))
            db.session.execute(db.text(f"INSERT OR IGNORE INTO filtros_maestros (proyecto_id, columna, valor) SELECT {pid}, columna, valor FROM filtros_maestros_old"))
            db.session.execute(db.text(f"INSERT OR IGNORE INTO tablas_maestras (proyecto_id, columna_criterio, valor_criterio, nueva_columna, nuevo_valor) SELECT {pid}, columna_criterio, valor_criterio, nueva_columna, nuevo_valor FROM tablas_maestras_old"))
        except Exception as e:
            print(f"Error restoring data: {e}")
        
        # Cleanup
        db.session.execute(db.text("DROP TABLE IF EXISTS nucleus_data_old"))
        db.session.execute(db.text("DROP TABLE IF EXISTS app_config_old"))
        db.session.execute(db.text("DROP TABLE IF EXISTS filtros_maestros_old"))
        db.session.execute(db.text("DROP TABLE IF EXISTS tablas_maestras_old"))
        db.session.commit()

    # Migration: Rename editor -> supervisor
    try:
        db.session.execute(db.text("UPDATE usuarios SET rol = 'supervisor' WHERE rol = 'editor'"))
        db.session.commit()
    except Exception:
        db.session.rollback()

    # Migration: Ensure fixed projects FLM, PEXT, Dataper, Material exist
    fixed = [('FLM', 'Fiscalización Lima Metropolitana'), ('PEXT', 'Proyecto Externo'), ('Dataper', 'DataPer S.A.C.'),
             ('Material', 'Materiales Disponibles')]
    for nombre, desc in fixed:
        if not Proyecto.query.filter_by(nombre=nombre).first():
            db.session.add(Proyecto(nombre=nombre, descripcion=desc))
    db.session.commit()

    # Migration: Configure Dataper columns (TECNICO, DOCUMENTO, CONTRATA, CELULAR, SUPERVISOR, DEPARTAMENTO, PROYECTO + ACTIVO/CESADO)
    dataper = Proyecto.query.filter_by(nombre='Dataper').first()
    if dataper:
        dataper_cols = [
            {'nombre': 'TECNICO', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'DOCUMENTO', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'CONTRATA', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'CELULAR', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'SUPERVISOR', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'DEPARTAMENTO', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'PROYECTO', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'ESTADO', 'tipo': 'lista', 'opciones': ['ACTIVO', 'CESADO']}
        ]
        mc_cfg = AppConfig.query.filter_by(proyecto_id=dataper.id, clave='manual_columns').first()
        if mc_cfg:
            mc_cfg.valor = json.dumps(dataper_cols, ensure_ascii=False)
        else:
            db.session.add(AppConfig(proyecto_id=dataper.id, clave='manual_columns', valor=json.dumps(dataper_cols, ensure_ascii=False)))

        pk_cfg = AppConfig.query.filter_by(proyecto_id=dataper.id, clave='primary_key').first()
        if not pk_cfg:
            db.session.add(AppConfig(proyecto_id=dataper.id, clave='primary_key', valor='DOCUMENTO'))

        schema_cfg = AppConfig.query.filter_by(proyecto_id=dataper.id, clave='app_schema').first()
        if not schema_cfg:
            db.session.add(AppConfig(proyecto_id=dataper.id, clave='app_schema', valor=json.dumps([])))
        db.session.commit()

    # Migration: Configure Material columns (COD_MATERIAL, DESCRIPCION_MATERIAL, PROYECTO, UM, TIPO)
    material_proy = Proyecto.query.filter_by(nombre='Material').first()
    if material_proy:
        material_proy.icono = 'fa-boxes-stacked'
        material_cols = [
            {'nombre': 'COD_MATERIAL', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'DESCRIPCION_MATERIAL', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'PROYECTO', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'UM', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'TIPO', 'tipo': 'texto', 'opciones': []}
        ]
        mc_cfg = AppConfig.query.filter_by(proyecto_id=material_proy.id, clave='manual_columns').first()
        if mc_cfg:
            mc_cfg.valor = json.dumps(material_cols, ensure_ascii=False)
        else:
            db.session.add(AppConfig(proyecto_id=material_proy.id, clave='manual_columns', valor=json.dumps(material_cols, ensure_ascii=False)))

        pk_cfg = AppConfig.query.filter_by(proyecto_id=material_proy.id, clave='primary_key').first()
        if pk_cfg:
            pk_cfg.valor = 'COD_MATERIAL'
        else:
            db.session.add(AppConfig(proyecto_id=material_proy.id, clave='primary_key', valor='COD_MATERIAL'))

        schema_cfg = AppConfig.query.filter_by(proyecto_id=material_proy.id, clave='app_schema').first()
        if not schema_cfg:
            db.session.add(AppConfig(proyecto_id=material_proy.id, clave='app_schema', valor=json.dumps([])))
        db.session.commit()

    # Migration: "Fault Level" es dato de origen (inmutable) -> no debe ser columna manual editable.
    # Se elimina de la configuracion para que nadie pueda editarla (ni admin).
    pext = Proyecto.query.filter_by(nombre='PEXT').first()
    if pext:
        mc_cfg = AppConfig.query.filter_by(proyecto_id=pext.id, clave='manual_columns').first()
        if mc_cfg:
            try:
                existing_cols = json.loads(mc_cfg.valor)
                if not isinstance(existing_cols, list): existing_cols = []
            except Exception:
                existing_cols = []
            # Remove the separate hours column (now shown inside Fault Level badge)
            # and Fault Level itself (source data, must not be editable)
            existing_cols = [c for c in existing_cols if c.get('nombre') not in ('Hrs Respuesta', 'Fault Level')]
            mc_cfg.valor = json.dumps(existing_cols, ensure_ascii=False)
        else:
            db.session.add(AppConfig(proyecto_id=pext.id, clave='manual_columns', valor=json.dumps([], ensure_ascii=False)))

        fault_rules = [
            ('Critical', '8hrs'),
            ('Alta', '10hrs'),
            ('Media', '48hrs'),
            ('Baja', '72hrs')
        ]
        for nivel, hrs in fault_rules:
            exists = TablaMaestra.query.filter_by(proyecto_id=pext.id, columna_criterio='Fault Level',
                                                  valor_criterio=nivel, nueva_columna='Hrs Respuesta').first()
            if not exists:
                db.session.add(TablaMaestra(proyecto_id=pext.id, columna_criterio='Fault Level',
                                            valor_criterio=nivel, nueva_columna='Hrs Respuesta', nuevo_valor=hrs))
        db.session.commit()

    # Migration: Seed SERVICIO options for WO detail (PEXT & FLM)
    default_servicios = ['PREVENTIVO', 'CORRECTIVO', 'PREDICTIVO', 'ABASTECIMIENTO DE COMBUSTIBLE',
                         'ADICIONALES', 'CORTE PROGRAMADO', 'TRABAJO PROGRAMADO']
    for proy_nombre in ('PEXT', 'FLM'):
        sp = Proyecto.query.filter_by(nombre=proy_nombre).first()
        if sp:
            scfg = AppConfig.query.filter_by(proyecto_id=sp.id, clave='servicio_opciones').first()
            if not scfg:
                db.session.add(AppConfig(proyecto_id=sp.id, clave='servicio_opciones',
                                         valor=json.dumps(default_servicios, ensure_ascii=False)))
    db.session.commit()

    # Migration: Campos del checklist de PEXT (datos del gestor al atender la incidencia)
    pext_checklist = Proyecto.query.filter_by(nombre='PEXT').first()
    if pext_checklist:
        mc_cfg = AppConfig.query.filter_by(proyecto_id=pext_checklist.id, clave='manual_columns').first()
        try:
            existing_cols = json.loads(mc_cfg.valor) if mc_cfg else []
            if not isinstance(existing_cols, list):
                existing_cols = []
        except Exception:
            existing_cols = []
        existing_names = {str(c.get('nombre', '')).strip() for c in existing_cols if isinstance(c, dict)}
        pext_new_cols = [
            {'nombre': 'MOTIVO PARADA RELOJ', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'ROBO HURTO', 'tipo': 'lista', 'opciones': ['Sí', 'No']},
            {'nombre': 'SUPERVISOR ATENCIÓN', 'tipo': 'texto', 'opciones': []},
            {'nombre': 'CORREO CIERRE', 'tipo': 'lista', 'opciones': ['Sí', 'No']},
            {'nombre': 'MOTIVO NO ATENDIDO', 'tipo': 'texto', 'opciones': []}
        ]
        for col in pext_new_cols:
            if col['nombre'] not in existing_names:
                existing_cols.append(col)
                existing_names.add(col['nombre'])
        if mc_cfg:
            mc_cfg.valor = json.dumps(existing_cols, ensure_ascii=False)
        else:
            db.session.add(AppConfig(proyecto_id=pext_checklist.id, clave='manual_columns',
                                     valor=json.dumps(existing_cols, ensure_ascii=False)))
        db.session.commit()

    # --- Backfill: historial de estado inicial para registros históricos ---
    # Los registros importados antes de existir el seguimiento de cambios de estado
    # no tienen fila en historial_cambios. Se crea una entrada 'IMPORT' con el estado
    # actual para que el historial/estado de cada WO sea visible y no se pierda info.
    # Es idempotente: se ejecuta una sola vez por proyecto (flag en app_config).
    WO_STATE_COL_BF = 'Estado de la tarea (WO State)'
    STATE_TS_COL_BF = 'FECHA CAMBIO ESTADO'
    for proy in Proyecto.query.all():
        bf_cfg = AppConfig.query.filter_by(proyecto_id=proy.id, clave='historial_backfill_done').first()
        if bf_cfg:
            continue
        hist_keys = set(kv for (kv,) in db.session.query(HistorialCambios.key_value).filter(
            HistorialCambios.proyecto_id == proy.id,
            HistorialCambios.campo_modificado == WO_STATE_COL_BF).distinct().all())
        n = 0
        for r in NucleusData.query.filter_by(proyecto_id=proy.id).all():
            if r.key_value in hist_keys:
                continue
            try:
                d = json.loads(r.data_json)
            except Exception:
                continue
            st = str(d.get(WO_STATE_COL_BF, '')).strip()
            if not st:
                continue
            fecha = datetime.utcnow()
            fec_txt = str(d.get(STATE_TS_COL_BF, '')).strip()
            if fec_txt:
                try:
                    fecha = datetime.strptime(fec_txt[:19], '%Y-%m-%d %H:%M:%S')
                except Exception:
                    pass
            db.session.add(HistorialCambios(
                proyecto_id=proy.id, usuario_id=None, username='IMPORT',
                key_value=r.key_value, campo_modificado=WO_STATE_COL_BF,
                valor_anterior='', valor_nuevo=st, fecha=fecha))
            n += 1
        if n:
            print(f"Backfill historial de estado: {n} registros (proyecto {proy.nombre})")
        db.session.add(AppConfig(proyecto_id=proy.id, clave='historial_backfill_done', valor='1'))
        db.session.commit()

def safe_json_dumps(obj):
    return json.dumps(obj, ensure_ascii=False)

# --- AUTH & PROJECT ROUTES ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = Usuario.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['rol'] = str(user.rol).strip().lower()
            
            # Default to first project with access if none active
            if 'current_proyecto_id' not in session:
                if user.rol in ['admin', 'demo']:
                    proj = Proyecto.query.first()
                else:
                    acceso = AccesoProyecto.query.filter_by(usuario_id=user.id).first()
                    proj = db.session.get(Proyecto, acceso.proyecto_id) if acceso else None
                
                if proj:
                    session['current_proyecto_id'] = int(proj.id)
                    session['current_proyecto_nombre'] = proj.nombre
                    
            return redirect(url_for('index'))
        return render_template('login.html', error="Credenciales inválidas")
    return render_template('login.html')

def get_session_info():
    uid = session.get('user_id')
    rol = str(session.get('rol') or 'supervisor').strip().lower()
    pid_raw = session.get('current_proyecto_id')
    pid = int(pid_raw) if pid_raw else None
    return uid, rol, pid

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/switch_project/<int:pid>')
@login_required
def switch_project(pid):
    # Check permission
    if session.get('rol') not in ['admin', 'demo']:
        acceso = AccesoProyecto.query.filter_by(usuario_id=session.get('user_id'), proyecto_id=pid).first()
        if not acceso:
            return redirect(url_for('index'))
            
    proj = db.session.get(Proyecto, pid)
    proj = db.session.get(Proyecto, pid)
    if proj:
        session['current_proyecto_id'] = int(proj.id)
        session['current_proyecto_nombre'] = proj.nombre
    
    # Redirect back to specified page, referrer, or index
    target = request.args.get('next') or request.referrer or url_for('index')
    return redirect(target)

# --- MAIN ROUTES ---
@app.route('/')
@login_required
def index():
    user_id, user_rol, pid = get_session_info()
    is_admin = user_rol == 'admin'
    is_privileged = user_rol in ['admin', 'demo']

    if not pid:
        # Emergency fallback or find first allowed
        if is_privileged:
            p = Proyecto.query.first()
        else:
            acc = AccesoProyecto.query.filter_by(usuario_id=user_id).first()
            p = db.session.get(Proyecto, acc.proyecto_id) if acc else None
            
        if p:
            session['current_proyecto_id'] = p.id
            session['current_proyecto_nombre'] = p.nombre
            pid = p.id
        else:
            session.clear()
            return render_template('login.html', error="No tiene proyectos asignados. Contacte al administrador.")

    # Verify access to current project
    res_obj = {}
    if not is_privileged:
        acc = AccesoProyecto.query.filter_by(usuario_id=user_id, proyecto_id=pid).first()
        if not acc:
            session.clear()
            return render_template('login.html', error="Acceso denegado a este proyecto. Por favor, solicite acceso al administrador.")
        try:
            res_obj = json.loads(acc.restricciones or '{}')
        except:
            res_obj = {}

    # Load all distinct keys reliably from AppConfig Master Schema
    schema_config = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
    columns_set = set(json.loads(schema_config.valor)) if schema_config else set()
    
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    manual_cols_data = json.loads(manual_cfg.valor) if manual_cfg else []
    for mc in manual_cols_data:
        columns_set.add(mc['nombre'])
        
    # Add KPI columns to the set so frontend can see them
    kpi_configs = KpiConfig.query.filter_by(proyecto_id=pid).all()
    for k in kpi_configs:
        if k.tipo == 'DILACION':
            columns_set.add(f"KPI_{k.nombre}")
    
    # Ocultar columnas internas (prefijo _) y redundantes de la vista
    columns_set = {c for c in columns_set if not c.startswith('_') and c != 'WO Number'}
    
    # Load and Filter data
    rows = NucleusData.query.filter_by(proyecto_id=pid).limit(2000).all()
    raw_data = []
    for r in rows:
        d = json.loads(r.data_json)
        d['_key'] = r.key_value
        raw_data.append(d)
        
    data = apply_data_restrictions(raw_data, res_obj)
        
    data, kpi_meta = inject_kpis(pid, data)

    config_key = AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').first()
    pk = config_key.valor if config_key else 'NO_DEF'
    
    # Keys con MÁS DE UN cambio de estado registrado (por importación o manual).
    # Cada transición de estado genera una fila en historial_cambios; el filtro
    # solo muestra los WO con más de una transición registrada (>= 2).
    WO_STATE_COL_CH = 'Estado de la tarea (WO State)'
    hist_counts = db.session.query(
        HistorialCambios.key_value,
        db.func.count(HistorialCambios.id)
    ).filter(
        HistorialCambios.proyecto_id == pid,
        HistorialCambios.campo_modificado == WO_STATE_COL_CH
    ).group_by(HistorialCambios.key_value).all()
    changed_keys_set = set(k for k, cnt in hist_counts if cnt > 1)
    changed_keys = sorted(changed_keys_set)
    
    cols = sorted(list(columns_set))
    if '_key' in cols: cols.remove('_key')
    cols.insert(0, '_key')
    
    # List allowed projects for the menu
    if is_privileged:
        proyectos = Proyecto.query.all()
    else:
        accesos = AccesoProyecto.query.filter_by(usuario_id=user_id).all()
        pids = [a.proyecto_id for a in accesos]
        proyectos = Proyecto.query.filter(Proyecto.id.in_(pids)).all()
    
    return render_template('index.html', 
                          data=json.dumps(data), 
                          columns=json.dumps(cols), 
                          pk=pk, 
                          manual_cols=json.dumps(manual_cols_data),
                          kpi_meta=json.dumps(kpi_meta),
                          changed_keys=json.dumps(changed_keys),
                          proyecto_id=pid,
                          proyectos_list=proyectos)

@app.route('/dashboard')
@login_required
def dashboard():
    user_id, user_rol, pid = get_session_info()
    is_admin = user_rol == 'admin'

    if not pid:
        return redirect(url_for('index'))

    # Verify access to current project
    res_obj = {}
    is_privileged = user_rol in ['admin', 'demo']

    if not is_privileged:
        acc = AccesoProyecto.query.filter_by(usuario_id=user_id, proyecto_id=pid).first()
        if not acc:
            session.clear()
            return render_template('login.html', error="Acceso denegado a este proyecto. Por favor, solicite acceso al administrador.")
        try:
            res_obj = json.loads(acc.restricciones or '{}')
        except:
            res_obj = {}

    schema_config = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
    columns_set = set(json.loads(schema_config.valor)) if schema_config else set()
    
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    manual_cols_data = json.loads(manual_cfg.valor) if manual_cfg else []
    for mc in manual_cols_data:
        columns_set.add(mc['nombre'])
        
    kpi_configs = KpiConfig.query.filter_by(proyecto_id=pid).all()
    for k in kpi_configs:
        if k.tipo == 'DILACION':
            columns_set.add(f"KPI_{k.nombre}")
    
    # Ocultar columnas internas (prefijo _) y redundantes de la vista
    columns_set = {c for c in columns_set if not c.startswith('_') and c != 'WO Number'}
    
    rows = NucleusData.query.filter_by(proyecto_id=pid).limit(5000).all()
    raw_data = []
    for r in rows:
        d = json.loads(r.data_json)
        d['_key'] = r.key_value
        raw_data.append(d)
        
    data = apply_data_restrictions(raw_data, res_obj)
        
    data, kpi_meta = inject_kpis(pid, data)

    cols = sorted(list(columns_set))
    
    # List allowed projects for the menu
    if is_privileged:
        proyectos = Proyecto.query.all()
    else:
        accesos = AccesoProyecto.query.filter_by(usuario_id=user_id).all()
        pids = [a.proyecto_id for a in accesos]
        proyectos = Proyecto.query.filter(Proyecto.id.in_(pids)).all()

    # Load saved configurations
    dash_config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_charts').first()
    saved_charts = json.loads(dash_config.valor) if dash_config else []
    
    kpi_config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_kpis').first()
    saved_kpis = json.loads(kpi_config.valor) if kpi_config else []
    
    filt_config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_filters').first()
    saved_filters = json.loads(filt_config.valor) if filt_config else []
    
    # Get current project name
    proj = Proyecto.query.get(pid)
    proyecto_nombre = proj.nombre if proj else "Gestión"
    
    return render_template('dashboard.html', 
                          data=json.dumps(data), 
                          columns=json.dumps(cols), 
                          proyectos_list=proyectos,
                          saved_charts=json.dumps(saved_charts),
                          saved_kpis=json.dumps(saved_kpis),
                          saved_filters=json.dumps(saved_filters),
                          proyecto_nombre=proyecto_nombre)

@app.route('/configuraciones')
@login_required
def configuraciones():
    user_rol = str(session.get('rol') or 'supervisor').strip().lower()
    if user_rol == 'gestor':
        return redirect(url_for('index'))
    user_id = session.get('user_id')
    if user_rol == 'admin':
        proyectos = Proyecto.query.all()
    else:
        accesos = AccesoProyecto.query.filter_by(usuario_id=user_id).all()
        pids = [a.proyecto_id for a in accesos]
        proyectos = Proyecto.query.filter(Proyecto.id.in_(pids)).all()
    return render_template('configuraciones.html', proyectos_list=proyectos)

@app.route('/admin')
@login_required
def admin_panel():
    if session.get('rol') != 'admin':
        return redirect(url_for('index'))
    return redirect(url_for('proyectos_page'))

@app.route('/proyectos')
@login_required
def proyectos_page():
    if session.get('rol') != 'admin':
        return redirect(url_for('index'))
    proy = Proyecto.query.order_by(Proyecto.id).all()
    return render_template('proyectos.html', proyectos=proy, proyectos_list=proy)

@app.route('/usuarios')
@login_required
def usuarios_page():
    if session.get('rol') != 'admin':
        return redirect(url_for('index'))
    proy = Proyecto.query.order_by(Proyecto.id).all()
    user = Usuario.query.all()
    accesos = {}
    for a in AccesoProyecto.query.all():
        accesos.setdefault(a.usuario_id, []).append(a.proyecto_id)
    return render_template('usuarios.html', proyectos=proy, usuarios=user, proyectos_list=proy, accesos=accesos)

# --- API ---
@app.route('/api/admin/proyecto', methods=['POST', 'DELETE'])
@login_required
def api_admin_proyecto():
    if session.get('rol') not in ['admin', 'demo']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar proyectos.'}), 403
    
    if request.method == 'POST':
        data = request.json
        nombre = data.get('nombre', '').strip()
        if not nombre: return jsonify({'error': 'Nombre requerido'}), 400
        try:
            nuevo = Proyecto(nombre=nombre, descripcion=data.get('descripcion', ''), icono=data.get('icono', 'fa-folder-open'))
            db.session.add(nuevo)
            db.session.commit()
            return jsonify({'success': True, 'id': nuevo.id})
        except:
            return jsonify({'error': 'Nombre duplicado'}), 400
            
    if request.method == 'DELETE':
        pid = request.json.get('id')
        if not pid: return jsonify({'error': 'ID requerido'}), 400
        p = db.session.get(Proyecto, pid)
        if p and p.nombre in ('FLM', 'PEXT', 'Dataper', 'Material'):
            return jsonify({'error': 'Los proyectos FLM, PEXT, Dataper y Material no se pueden eliminar.'}), 403
        try:
            # Cascading delete manually for safety (or set up models with cascade)
            # We must not delete the project 1 (Pangeaco) if it's the only one or a protected one?
            # User choice, I'll allow deleting any.
            Tecnico.query.filter_by(proyecto_id=pid).delete()
            NucleusData.query.filter_by(proyecto_id=pid).delete()
            AppConfig.query.filter_by(proyecto_id=pid).delete()
            FiltroMaestro.query.filter_by(proyecto_id=pid).delete()
            TablaMaestra.query.filter_by(proyecto_id=pid).delete()
            
            p = db.session.get(Proyecto, pid)
            if p:
                db.session.delete(p)
                db.session.commit()
                # If deleted project is active project, clear it
                if session.get('current_proyecto_id') == int(pid):
                    session.pop('current_proyecto_id', None)
                    session.pop('current_proyecto_nombre', None)
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

@app.route('/api/tecnicos', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_tecnicos():
    if session.get('rol') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 403

    if request.method == 'GET':
        pid = request.args.get('proyecto_id', type=int)
        q = Tecnico.query
        if pid:
            q = q.filter_by(proyecto_id=pid)
        return jsonify([{
            'id': t.id,
            'proyecto_id': t.proyecto_id,
            'nombre': t.nombre,
            'contrata': t.contrata,
            'especialidad': t.especialidad,
            'telefono': t.telefono
        } for t in q.order_by(Tecnico.nombre).all()])

    if request.method == 'POST':
        data = request.json
        pid = data.get('proyecto_id')
        nombre = (data.get('nombre') or '').strip()
        if not pid or not nombre:
            return jsonify({'error': 'Proyecto y nombre requeridos'}), 400
        nuevo = Tecnico(
            proyecto_id=pid,
            nombre=nombre,
            contrata=(data.get('contrata') or '').strip(),
            especialidad=(data.get('especialidad') or '').strip(),
            telefono=(data.get('telefono') or '').strip()
        )
        db.session.add(nuevo)
        db.session.commit()
        return jsonify({'success': True, 'id': nuevo.id})

    if request.method == 'DELETE':
        tid = request.json.get('id')
        if not tid:
            return jsonify({'error': 'ID requerido'}), 400
        t = db.session.get(Tecnico, tid)
        if not t:
            return jsonify({'error': 'No encontrado'}), 404
        db.session.delete(t)
        db.session.commit()
        return jsonify({'success': True})

@app.route('/api/admin/usuario', methods=['POST', 'PUT', 'DELETE'])
@login_required
def api_admin_usuario():
    if session.get('rol') not in ['admin', 'demo']:
        return jsonify({'error': 'Unauthorized'}), 403
        
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar usuarios.'}), 403
        
    if request.method == 'POST':
        data = request.json
        nombre = data.get('nombre', '').strip()
        user = data.get('username', '').strip()
        pw = data.get('password', '').strip()
        rol = data.get('rol', 'supervisor').strip()
        proyectos = data.get('proyectos', [])
        if not all([user, pw]): return jsonify({'error': 'Datos incompletos'}), 400
        try:
            nuevo = Usuario(username=user, password_hash=generate_password_hash(pw), rol=rol, nombre=nombre)
            db.session.add(nuevo)
            db.session.flush()
            for pid in proyectos:
                db.session.add(AccesoProyecto(usuario_id=nuevo.id, proyecto_id=int(pid), restricciones='{}'))
            db.session.commit()
            return jsonify({'success': True, 'id': nuevo.id})
        except:
            db.session.rollback()
            return jsonify({'error': 'Usuario duplicado'}), 400
            
    if request.method == 'PUT':
        data = request.json
        uid = data.get('id')
        nombre = data.get('nombre', '').strip()
        user = data.get('username', '').strip()
        pw = data.get('password', '').strip()
        rol = data.get('rol', '').strip()
        proyectos = data.get('proyectos')
        
        if not uid or not user: return jsonify({'error': 'ID y usuario requeridos'}), 400
        try:
            u = db.session.get(Usuario, uid)
            if not u: return jsonify({'error': 'Usuario no encontrado'}), 404
            
            u.nombre = nombre
            u.username = user
            if pw:
                u.password_hash = generate_password_hash(pw)
            if rol:
                u.rol = rol
                
            if proyectos is not None:
                AccesoProyecto.query.filter_by(usuario_id=uid).delete()
                for pid in proyectos:
                    db.session.add(AccesoProyecto(usuario_id=uid, proyecto_id=int(pid), restricciones='{}'))
                
            db.session.commit()
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': 'Usuario duplicado o error: ' + str(e)}), 400
            
    if request.method == 'DELETE':
        uid = request.json.get('id')
        if not uid: return jsonify({'error': 'ID requerido'}), 400
        if int(uid) == session.get('user_id'):
            return jsonify({'error': 'No puedes borrar tu propio usuario'}), 400
        try:
            u = db.session.get(Usuario, uid)
            if u:
                # Delete permissions too
                AccesoProyecto.query.filter_by(usuario_id=uid).delete()
                db.session.delete(u)
                db.session.commit()
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

@app.route('/api/admin/usuario/duplicar', methods=['POST'])
@login_required
def api_admin_usuario_duplicar():
    if session.get('rol') not in ['admin']:
        return jsonify({'error': 'Solo admin puede duplicar usuarios.'}), 403
        
    data = request.json
    src_id = data.get('source_id')
    new_user = data.get('new_username', '').strip()
    new_pw = data.get('new_password', '').strip()
    
    if not all([src_id, new_user, new_pw]):
        return jsonify({'error': 'Datos incompletos para duplicar'}), 400
        
    src_u = db.session.get(Usuario, src_id)
    if not src_u: 
        return jsonify({'error': 'Usuario origen no encontrado'}), 404
        
    try:
        nuevo = Usuario(username=new_user, password_hash=generate_password_hash(new_pw), rol=src_u.rol)
        db.session.add(nuevo)
        db.session.flush() # Para obtener el nuevo ID
        
        # Copiar permisos (AccesoProyecto)
        permisos = AccesoProyecto.query.filter_by(usuario_id=src_id).all()
        for p in permisos:
            new_p = AccesoProyecto(usuario_id=nuevo.id, proyecto_id=p.proyecto_id, restricciones=p.restricciones)
            db.session.add(new_p)
            
        db.session.commit()
        return jsonify({'success': True, 'id': nuevo.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': 'Error al duplicar (¿Usuario duplicado?): ' + str(e)}), 400

@app.route('/api/admin/permisos', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_admin_permisos():
    if session.get('rol') not in ['admin', 'demo']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar permisos.'}), 403
    
    if request.method == 'GET':
        uid = request.args.get('uid')
        if not uid: return jsonify([])
        permisos = AccesoProyecto.query.filter_by(usuario_id=uid).all()
        result = []
        for p in permisos:
            proj = db.session.get(Proyecto, p.proyecto_id)
            result.append({
                'id': p.id,
                'proyecto_id': p.proyecto_id,
                'proyecto_nombre': proj.nombre if proj else 'Desconocido',
                'restricciones': p.restricciones
            })
        return jsonify(result)

    if request.method == 'POST':
        data = request.json
        uid = data.get('usuario_id')
        pid = data.get('proyecto_id')
        res = data.get('restricciones', '{}')
        if not uid or not pid: return jsonify({'error': 'Faltan datos'}), 400
        
        existente = AccesoProyecto.query.filter_by(usuario_id=uid, proyecto_id=pid).first()
        if existente:
            existente.restricciones = res
        else:
            nuevo = AccesoProyecto(usuario_id=uid, proyecto_id=pid, restricciones=res)
            db.session.add(nuevo)
        db.session.commit()
        return jsonify({'success': True})

    if request.method == 'DELETE':
        aid = request.json.get('id')
        acc = db.session.get(AccesoProyecto, aid)
        if acc:
            db.session.delete(acc)
            db.session.commit()
        return jsonify({'success': True})

@app.route('/api/admin/columnas')
@login_required
def api_admin_columnas():
    if session.get('rol') not in ['admin', 'demo']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    pid = request.args.get('pid')
    if not pid: return jsonify([])
    
    # Standard columns
    schema_config = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
    cols = json.loads(schema_config.valor) if schema_config else []
    
    # Manual columns
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    if manual_cfg:
        m_cols = json.loads(manual_cfg.valor)
        for mc in m_cols:
            if mc['nombre'] not in cols:
                cols.append(mc['nombre'])
                
    return jsonify(sorted(cols))

@app.route('/api/import/manual_template')
@login_required
def api_import_manual_template():
    """Generates and downloads an Excel template with primary key + manual columns."""
    from flask import make_response
    pid = session.get('current_proyecto_id')
    if not pid:
        return jsonify({'error': 'No project selected'}), 400

    # Get the primary key column name
    config_key = AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').first()
    pk_name = config_key.valor if config_key else '_key'

    # Get manual columns
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    manual_cols = [mc['nombre'] for mc in json.loads(manual_cfg.valor)] if manual_cfg else []

    if not manual_cols:
        return jsonify({'error': 'No hay columnas manuales configuradas en este proyecto.'}), 400

    # Fetch all existing primary key values
    rows = NucleusData.query.filter_by(proyecto_id=pid).all()

    proy = db.session.get(Proyecto, pid)
    proy_nombre = proy.nombre if proy else ''

    # Plantilla del checklist de PEXT: PK + columnas que cubren el checklist.
    # Las columnas importadas (que ya tienen dato) van con su valor actual;
    # las manuales van vacías para que el gestor las llene.
    PEXT_CHECKLIST = [
        'Causa raíz',
        'Nombre de Site',
        'Estado del TT',
        'Fecha de creación (WO Creation date)',
        'Fecha y hora de WO a estado close',
        'Fault Level',
    ]
    PEXT_CHECKLIST_MANUAL = {
        'MATERIAL USADO', 'INICIO DE PARADA', 'FIN DE PARADA',
        'REQUIERE CORRECTIVO FINAL', 'SOLUCIÓN',
        'MOTIVO PARADA RELOJ', 'ROBO HURTO', 'SUPERVISOR ATENCIÓN',
        'CORREO CIERRE', 'MOTIVO NO ATENDIDO'
    }
    if proy_nombre == 'PEXT':
        checklist_cols = PEXT_CHECKLIST + [c for c in manual_cols if c in PEXT_CHECKLIST_MANUAL]
        data = {}
        for r in rows:
            d = json.loads(r.data_json)
            line = {pk_name: r.key_value}
            for col in PEXT_CHECKLIST:
                line[col] = str(d.get(col, '') or '')
            for col in checklist_cols:
                if col not in PEXT_CHECKLIST:
                    line[col] = ''
            data[r.key_value] = line
        if data:
            df = pd.DataFrame(list(data.values()))
        else:
            df = pd.DataFrame(columns=[pk_name] + checklist_cols)
        # Auto-adjust column widths
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Plantilla')
            ws = writer.sheets['Plantilla']
            for col_cells in ws.columns:
                max_len = max(len(str(cell.value or '')) for cell in col_cells)
                ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 45)
        output.seek(0)
        response = make_response(output.read())
        response.headers['Content-Disposition'] = 'attachment; filename=plantilla_checklist_PEXT.xlsx'
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        return response

    # Build DataFrame: pk column pre-filled, manual columns empty
    df_data = {pk_name: [r.key_value for r in rows]}
    for col in manual_cols:
        df_data[col] = [''] * len(rows)
    df = pd.DataFrame(df_data)

    # Write to in-memory Excel
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Plantilla')
        # Auto-adjust column widths
        ws = writer.sheets['Plantilla']
        for col_cells in ws.columns:
            max_len = max(len(str(cell.value or '')) for cell in col_cells)
            ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 40)
    output.seek(0)

    response = make_response(output.read())
    response.headers['Content-Disposition'] = 'attachment; filename=plantilla_columnas_manuales.xlsx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    return response

@app.route('/api/import/preview', methods=['POST'])
@login_required
def api_import_preview():
    file = request.files.get('file')
    if not file: return jsonify({'error': 'No file'}), 400
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file, encoding='utf-8', nrows=5, dtype=str)
        else:
            df = pd.read_excel(file, nrows=5, dtype=str)
        return jsonify({'columns': [str(c).strip() for c in df.columns]})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/import/process', methods=['POST'])
@login_required
def api_import_process():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para realizar importaciones.'}), 403
    import_type = request.form.get('type') or 'base' # 'base' or 'cruce'
    sum_duplicates = request.form.get('sum_duplicates') == 'true'
    sum_type = request.form.get('sum_type', 'number')
    consolidate_date = request.form.get('consolidate_date') == 'true'
    date_column = request.form.get('date_column', '').strip()
    cols_to_keep_str = request.form.get('columns_to_keep', '[]')
    columns_to_keep = json.loads(cols_to_keep_str)
    file_key = request.form.get('file_key', '').strip()

    # Columnas manuales del proyecto (datos editados por gestores)
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    manual_columns_list = []
    if manual_cfg:
        for mc in json.loads(manual_cfg.valor):
            if isinstance(mc, dict):
                manual_columns_list.append(str(mc.get('nombre', '')).strip())
    manual_columns_list = [c for c in manual_columns_list if c]
    
    file = request.files.get('file')
    if not file: return jsonify({'error': 'No file'}), 400
    
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file, encoding='utf-8', dtype=str)
        else:
            df = pd.read_excel(file, dtype=str)
            
        df.columns = [str(c).strip() for c in df.columns]
        if not file_key:
            pk_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').first()
            candidate = (pk_cfg.valor if pk_cfg else None) or ''
            file_key = candidate if candidate in df.columns else str(df.columns[0])
        if file_key not in df.columns:
            return jsonify({'error': f'Key {file_key} not found in headers.'}), 400

        # Modo columnas manuales: por defecto solo llave + columnas manuales del archivo
        if import_type == 'manual_cols' and not columns_to_keep:
            columns_to_keep = [file_key] + [c for c in manual_columns_list if c in df.columns and c != file_key]

        if columns_to_keep:
            valid_cols = [c for c in columns_to_keep if c in df.columns]
            if file_key not in valid_cols: valid_cols.append(file_key)
            if consolidate_date and date_column and date_column in df.columns and date_column not in valid_cols:
                valid_cols.append(date_column)
            df = df[valid_cols]
        df = df.fillna('')

        if consolidate_date and date_column and date_column in df.columns and len(df) > 0:
            df['_temp_date'] = pd.to_datetime(df[date_column], errors='coerce')
            df = df.sort_values(by='_temp_date', ascending=True, na_position='first')
            df = df.drop_duplicates(subset=[file_key], keep='last')
            df = df.drop(columns=['_temp_date'])
        elif sum_duplicates and len(df) > 0:
            # Smart aggregation: Sum numeric columns, last for the rest
            agg_dict = {}
            for col in df.columns:
                if col == file_key: continue
                
                if sum_type == 'soles':
                    # Clean currency formatting before numeric conversion
                    cleaned_col = df[col].astype(str).str.replace(r'[sS]/\.?\s*', '', regex=True).str.replace(',', '')
                    temp_numeric = pd.to_numeric(cleaned_col, errors='coerce')
                else:
                    temp_numeric = pd.to_numeric(df[col], errors='coerce')
                    
                if not temp_numeric.isna().all():
                    df[col] = temp_numeric.fillna(0)
                    agg_dict[col] = 'sum'
                else:
                    agg_dict[col] = 'last'
            
            if agg_dict:
                df = df.groupby(file_key, as_index=False).agg(agg_dict)
        
        filtros = FiltroMaestro.query.filter_by(proyecto_id=pid).all()
        
        # Agrupar reglas por "Clusters" de columnas (Connected Components)
        # Esto permite que reglas para diferentes columnas se sumen con AND
        # Pero reglas que comparten columnas se sumen con OR
        raw_reglas = []
        for f in filtros:
            cols = [c.strip() for c in f.columna.split(',')]
            vals = [v.strip() for v in f.valor.split(',')]
            raw_reglas.append({'cols': set(cols), 'pairs': list(zip(cols, vals))})
            
        clusters = []
        for r in raw_reglas:
            assigned = False
            for group in clusters:
                # Si la regla comparte alguna columna con el grupo, se une a él
                if any(c in group['columns'] for c in r['cols']):
                    group['columns'].update(r['cols'])
                    group['rules'].append(r['pairs'])
                    assigned = True
                    break
            if not assigned:
                clusters.append({'columns': r['cols'], 'rules': [r['pairs']]})
        
        # Consolidar clusters que puedan haberse cruzado después de unirse por partes
        final_clusters = []
        for c in clusters:
            merged = False
            for f in final_clusters:
                if c['columns'] & f['columns']:
                    f['columns'].update(c['columns'])
                    f['rules'].extend(c['rules'])
                    merged = True
                    break
            if not merged:
                final_clusters.append(c)
            
        tablas = TablaMaestra.query.filter_by(proyecto_id=pid).all()
        reglas = []
        for t in tablas:
            t_cols = [c.strip() for c in t.columna_criterio.split(',')]
            t_vals = [v.strip() for v in t.valor_criterio.split(',')]
            reglas.append({
                'condiciones': list(zip(t_cols, t_vals)),
                'nueva_columna': t.nueva_columna,
                'nuevo_valor': t.nuevo_valor
            })

        reglas_manuales = ReglaEstadoManual.query.filter_by(proyecto_id=pid).all()
        for r in reglas_manuales:
            r_cols = [c.strip() for c in r.columna_criterio.split(',')]
            r_vals = [v.strip() for v in r.valor_criterio.split(',')]
            reglas.append({
                'condiciones': list(zip(r_cols, r_vals)),
                'nueva_columna': r.columna_manual,
                'nuevo_valor': r.nuevo_valor
            })

        config_pk = AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').first()
        system_pk = config_pk.valor if config_pk else None
        
        if import_type == 'base' and not system_pk:
            system_pk = file_key
            db.session.add(AppConfig(proyecto_id=pid, clave='primary_key', valor=system_pk))
            db.session.commit()
            
        existing_records_list = NucleusData.query.filter_by(proyecto_id=pid).all()
        existing_records = {r.key_value: r for r in existing_records_list}
        
        # Consolidation Config
        cons_cfg_row = AppConfig.query.filter_by(proyecto_id=pid, clave='consolidation_config').first()
        cons_cfg = json.loads(cons_cfg_row.valor) if cons_cfg_row else {}
        consolidate_on_fail = cons_cfg.get('consolidate_on_filter_fail', False)
        
        updated, added, ignored, consolidated = 0, 0, 0, 0
        dynamic_cols = set()
        imported_keys = set()

        WO_STATE_COL = 'Estado de la tarea (WO State)'
        STATE_TS_COL = 'FECHA CAMBIO ESTADO'
        DISPATCHED_TS_COL = '_fecha_dispatched'
        CANCEL_REJECT_TS_COL = '_fecha_cancel_reject'
        DISPATCHED_STATES = {'dispatched'}
        TERMINAL_STATES = {'canceled', 'rejected'}
        now_str = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')

        # Campos protegidos: la importación NO pisa los datos editados manualmente.
        # EXCEPCIÓN: en modo 'manual_cols' (Dataper) el archivo ES la fuente de datos manuales.
        protected_fields = set(manual_columns_list)
        protected_fields.update([
            'SERVICIO', 'CIUDAD', 'TECNICO', 'CONTRATA',
            'MOTIVO DE AVERÍA', 'MOTIVO DE AVERIA', 'SOLUCIÓN', 'SOLUCION',
            'LATITUD', 'LONGITUD', 'SE INSTALÓ MUFAS', 'SE INSTALO MUFAS',
            'LATITUD MUFAS', 'LATITUD Mufas', 'LONGITUD MUFAS', 'LONGITUD Mufas',
            'UBICACIÓN DE MUFAS', 'UBICACION DE MUFAS',
            'MATERIALES'
        ])
        
        for idx, row in df.iterrows():
            row_dict = row.to_dict()
            key_val = str(row_dict.get(file_key, '')).strip()
            if not key_val: continue
            imported_keys.add(key_val)
            
            # --- 1. MERGE CON DATA EXISTENTE ---
            is_new = False
            if key_val in existing_records:
                record = existing_records[key_val]
                current_data = json.loads(record.data_json)
                old_state = str(current_data.get(WO_STATE_COL, '')).strip()
                if import_type == 'manual_cols':
                    current_data.update(row_dict)
                else:
                    for k, v in row_dict.items():
                        if k in protected_fields and k in current_data:
                            continue
                        current_data[k] = v
                new_state = str(current_data.get(WO_STATE_COL, '')).strip()
                new_state_l = new_state.lower()
                if new_state != old_state:
                    # Registrar TODA transición (incluye quedarse vacío), de modo
                    # que el historial nunca pierda cambios de estado por importación.
                    if new_state_l != old_state.lower():
                        current_data[STATE_TS_COL] = now_str
                        dynamic_cols.add(STATE_TS_COL)
                    db.session.add(HistorialCambios(
                        proyecto_id=pid,
                        usuario_id=None,
                        username='IMPORT',
                        key_value=key_val,
                        campo_modificado=WO_STATE_COL,
                        valor_anterior=old_state,
                        valor_nuevo=new_state,
                        fecha=datetime.utcnow()
                    ))
                if new_state_l in DISPATCHED_STATES and not current_data.get(DISPATCHED_TS_COL):
                    current_data[DISPATCHED_TS_COL] = now_str
                if new_state_l in TERMINAL_STATES and not current_data.get(CANCEL_REJECT_TS_COL):
                    current_data[CANCEL_REJECT_TS_COL] = now_str
            else:
                if import_type == 'cruce':
                    schema_config = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
                    schema_val = json.loads(schema_config.valor) if schema_config else []
                    if schema_val:
                        continue
                current_data = row_dict.copy()
                init_state = str(current_data.get(WO_STATE_COL, '')).strip().lower()
                if init_state:
                    current_data[STATE_TS_COL] = now_str
                    dynamic_cols.add(STATE_TS_COL)
                    db.session.add(HistorialCambios(
                        proyecto_id=pid,
                        usuario_id=None,
                        username='IMPORT',
                        key_value=key_val,
                        campo_modificado=WO_STATE_COL,
                        valor_anterior='',
                        valor_nuevo=str(current_data.get(WO_STATE_COL, '')).strip(),
                        fecha=datetime.utcnow()
                    ))
                if init_state in DISPATCHED_STATES:
                    current_data[DISPATCHED_TS_COL] = now_str
                if init_state in TERMINAL_STATES:
                    current_data[CANCEL_REJECT_TS_COL] = now_str
                is_new = True

            # --- 2. APLICAR REGLAS Y TABLAS MAESTRAS ---
            for regla in reglas:
                match = True
                for c, v in regla['condiciones']:
                    if str(current_data.get(c, '')) != v:
                        match = False; break
                if match:
                    current_data[regla['nueva_columna']] = regla['nuevo_valor']
                    dynamic_cols.add(regla['nueva_columna'])

            # --- 3. LOGICA DE FILTRO POR CLUSTERS ---
            keep_record = True
            if final_clusters:
                for cluster in final_clusters:
                    # Usamos current_data para evaluar. Si es un cruce, ahora tiene la info base también.
                    match_cluster = False
                    for rule in cluster['rules']:
                        match_rule = True
                        for c, v in rule:
                            val_archivo = str(current_data.get(c, '')).strip().upper()
                            val_filtro = str(v).strip().upper()
                            if val_archivo != val_filtro:
                                match_rule = False
                                break
                        if match_rule:
                            match_cluster = True
                            break
                    if not match_cluster:
                        keep_record = False
                        break
            
            # --- 4. DECISIÓN DE CONSOLIDACIÓN ---
            if not keep_record:
                if consolidate_on_fail and key_val in existing_records:
                    record = existing_records[key_val]
                    # Guardamos el current_data (que contiene las actualizaciones del archivo importado) en el histórico
                    new_hist = NucleusHistory(proyecto_id=pid, key_value=key_val, data_json=safe_json_dumps(current_data))
                    db.session.add(new_hist)
                    db.session.delete(record)
                    consolidated += 1
                    updated += 1
                    del existing_records[key_val]
                else:
                    ignored += 1
                continue
                
            # --- 5. GUARDAR REGISTRO ACTIVO ---
            if is_new:
                new_record = NucleusData(proyecto_id=pid, key_value=key_val, data_json=safe_json_dumps(current_data))
                db.session.add(new_record)
                existing_records[key_val] = new_record
                added += 1
            else:
                record = existing_records[key_val]
                record.data_json = safe_json_dumps(current_data)
                updated += 1
        
        # Absence-based Consolidation (Optimized to avoid SQLite parameter limits)
        absent_consolidated = 0
        if import_type == 'base' and cons_cfg.get('auto_consolidate_missing'):
            # Python-based comparison is safer and faster for large datasets
            # existing_records contains all records of the project currently in DB
            for kv, rec in existing_records.items():
                if kv not in imported_keys:
                    new_hist = NucleusHistory(proyecto_id=pid, key_value=rec.key_value, data_json=rec.data_json)
                    db.session.add(new_hist)
                    db.session.delete(rec)
                    absent_consolidated += 1
                    consolidated += 1

        db.session.commit()
        
        config_schema = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
        schema_cols = set(json.loads(config_schema.valor)) if config_schema else set()
        new_schema = schema_cols.union(set(df.columns)).union(dynamic_cols)
        if new_schema != schema_cols:
            if config_schema:
                config_schema.valor = safe_json_dumps(list(new_schema))
            else:
                db.session.add(AppConfig(proyecto_id=pid, clave='app_schema', valor=safe_json_dumps(list(new_schema))))
            db.session.commit()
        return jsonify({
            'success': True, 
            'added': added, 
            'updated': updated, 
            'ignored': ignored, 
            'consolidated': consolidated,
            'pk': system_pk
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/master/filtros', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_master_filtros():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar filtros.'}), 403
    if request.method == 'GET':
        qs = FiltroMaestro.query.filter_by(proyecto_id=pid).all()
        return jsonify([{'id': q.id, 'columna': q.columna, 'valor': q.valor} for q in qs])
    if request.method == 'POST':
        data = request.json
        try:
            nuevo = FiltroMaestro(proyecto_id=pid, columna=data['columna'].strip(), valor=data['valor'].strip())
            db.session.add(nuevo)
            db.session.commit()
            return jsonify({'success': True, 'id': nuevo.id})
        except:
            db.session.rollback()
            return jsonify({'error': 'Duplicated or invalid'}), 400
    if request.method == 'DELETE':
        data = request.json
        if data.get('clear_all'):
            FiltroMaestro.query.filter_by(proyecto_id=pid).delete()
            db.session.commit()
            return jsonify({'success': True})
        
        id = data.get('id')
        if id:
            f = FiltroMaestro.query.filter_by(id=id, proyecto_id=pid).first()
            if f:
                db.session.delete(f)
                db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/tablas', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_master_tablas():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar tablas maestras.'}), 403
    if request.method == 'GET':
        qs = TablaMaestra.query.filter_by(proyecto_id=pid).all()
        return jsonify([{'id': q.id, 'columna_criterio': q.columna_criterio, 'valor_criterio': q.valor_criterio, 'nueva_columna': q.nueva_columna, 'nuevo_valor': q.nuevo_valor} for q in qs])
    if request.method == 'POST':
        data = request.json
        nuevo = TablaMaestra(
            proyecto_id=pid,
            columna_criterio=data['columna_criterio'].strip(),
            valor_criterio=data['valor_criterio'].strip(),
            nueva_columna=data['nueva_columna'].strip(),
            nuevo_valor=data['nuevo_valor'].strip(),
        )
        db.session.add(nuevo)
        db.session.commit()
        return jsonify({'success': True, 'id': nuevo.id})
    if request.method == 'DELETE':
        data = request.json
        if data.get('clear_all'):
            TablaMaestra.query.filter_by(proyecto_id=pid).delete()
            db.session.commit()
            return jsonify({'success': True})
            
        id = data.get('id')
        if id:
            f = TablaMaestra.query.filter_by(id=id, proyecto_id=pid).first()
            if f:
                db.session.delete(f)
                db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/reglas_manuales', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_master_reglas_manuales():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar reglas manuales.'}), 403
    if request.method == 'GET':
        qs = ReglaEstadoManual.query.filter_by(proyecto_id=pid).all()
        return jsonify([{'id': q.id, 'columna_criterio': q.columna_criterio, 'valor_criterio': q.valor_criterio, 'columna_manual': q.columna_manual, 'nuevo_valor': q.nuevo_valor} for q in qs])
    if request.method == 'POST':
        data = request.json
        nuevo = ReglaEstadoManual(
            proyecto_id=pid,
            columna_criterio=data['columna_criterio'].strip(),
            valor_criterio=data['valor_criterio'].strip(),
            columna_manual=data['columna_manual'].strip(),
            nuevo_valor=data['nuevo_valor'].strip(),
        )
        db.session.add(nuevo)
        db.session.commit()
        return jsonify({'success': True, 'id': nuevo.id})
    if request.method == 'DELETE':
        data = request.json
        if data.get('clear_all'):
            ReglaEstadoManual.query.filter_by(proyecto_id=pid).delete()
            db.session.commit()
            return jsonify({'success': True})
            
        id = data.get('id')
        if id:
            f = ReglaEstadoManual.query.filter_by(id=id, proyecto_id=pid).first()
            if f:
                db.session.delete(f)
                db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/reprocess', methods=['POST'])
@login_required
def api_master_reprocess():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para re-procesar datos.'}), 403
    try:
        # 1. Load data update rules (Tablas Maestras & Reglas Manuales)
        tablas = TablaMaestra.query.filter_by(proyecto_id=pid).all()
        reglas = []
        for t in tablas:
            t_cols = [c.strip() for c in t.columna_criterio.split(',')]
            t_vals = [v.strip() for v in t.valor_criterio.split(',')]
            reglas.append({
                'condiciones': list(zip(t_cols, t_vals)),
                'nueva_columna': t.nueva_columna,
                'nuevo_valor': t.nuevo_valor
            })
            
        reglas_manuales = ReglaEstadoManual.query.filter_by(proyecto_id=pid).all()
        for r in reglas_manuales:
            r_cols = [c.strip() for c in r.columna_criterio.split(',')]
            r_vals = [v.strip() for v in r.valor_criterio.split(',')]
            reglas.append({
                'condiciones': list(zip(r_cols, r_vals)),
                'nueva_columna': r.columna_manual,
                'nuevo_valor': r.nuevo_valor
            })

        # 2. Load Filter Rules (Clusters)
        filtros = FiltroMaestro.query.filter_by(proyecto_id=pid).all()
        raw_reglas_filtros = []
        for f in filtros:
            f_cols = [c.strip() for c in f.columna.split(',')]
            f_vals = [v.strip() for v in f.valor.split(',')]
            raw_reglas_filtros.append({'cols': set(f_cols), 'pairs': list(zip(f_cols, f_vals))})
            
        temp_clusters = []
        for r in raw_reglas_filtros:
            assigned = False
            for group in temp_clusters:
                if any(c in group['columns'] for c in r['cols']):
                    group['columns'].update(r['cols'])
                    group['rules'].append(r['pairs'])
                    assigned = True
                    break
            if not assigned:
                temp_clusters.append({'columns': r['cols'], 'rules': [r['pairs']]})
                
        # Consolidar clusters transitivos
        final_clusters = []
        for c in temp_clusters:
            merged = False
            for f in final_clusters:
                if c['columns'] & f['columns']:
                    f['columns'].update(c['columns'])
                    f['rules'].extend(c['rules'])
                    merged = True
                    break
            if not merged:
                final_clusters.append(c)

        # 3. Load Consolidation Config
        cons_cfg_row = AppConfig.query.filter_by(proyecto_id=pid, clave='consolidation_config').first()
        cons_cfg = json.loads(cons_cfg_row.valor) if cons_cfg_row else {}
        consolidate_on_fail = cons_cfg.get('consolidate_on_filter_fail', False)

        # 4. Process Data
        records = NucleusData.query.filter_by(proyecto_id=pid).all()
        updated, consolidated = 0, 0
        new_columns = set()
        
        for record in records:
            row_dict = json.loads(record.data_json)
            data_changed = False
            
            # Apply update rules
            for regla in reglas:
                match = True
                for c, v in regla['condiciones']:
                    if str(row_dict.get(c, '')) != v:
                        match = False; break
                if match:
                    if row_dict.get(regla['nueva_columna']) != regla['nuevo_valor']:
                        row_dict[regla['nueva_columna']] = regla['nuevo_valor']
                        new_columns.add(regla['nueva_columna'])
                        data_changed = True
            
            # Check Consolidation Rules
            move_to_history = False
            
            # Filter-fail-based
            if consolidate_on_fail and final_clusters:
                keep_record = True
                for cluster in final_clusters:
                    match_cluster = False
                    for rule in cluster['rules']:
                        match_rule = True
                        for c, v in rule:
                            val_archivo = str(row_dict.get(c, '')).strip().upper()
                            val_filtro = str(v).strip().upper()
                            if val_archivo != val_filtro:
                                match_rule = False
                                break
                        if match_rule:
                            match_cluster = True
                            break
                    if not match_cluster:
                        keep_record = False
                        break
                
                if not keep_record:
                    move_to_history = True

            if move_to_history:
                new_hist = NucleusHistory(proyecto_id=pid, key_value=record.key_value, data_json=safe_json_dumps(row_dict))
                db.session.add(new_hist)
                db.session.delete(record)
                consolidated += 1
                updated += 1
            elif data_changed:
                record.data_json = safe_json_dumps(row_dict)
                updated += 1
        
        db.session.commit()
        
        # Update schema if new columns were found
        if new_columns:
            config_schema = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
            if config_schema:
                schema_cols = set(json.loads(config_schema.valor))
                if not new_columns.issubset(schema_cols):
                    updated_schema = list(schema_cols.union(new_columns))
                    config_schema.valor = safe_json_dumps(updated_schema)
                    db.session.commit()

        return jsonify({'success': True, 'updated': updated, 'consolidated': consolidated})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/master/manual_columns', methods=['GET', 'POST', 'DELETE'])
@login_required
def api_manual_columns():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar columnas manuales.'}), 403
    config = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    cols = json.loads(config.valor) if config else []
    if request.method == 'GET':
        return jsonify(cols)
    if request.method == 'POST':
        data = request.json
        new_col = {
            'nombre': data['nombre'].strip(),
            'tipo': data['tipo'].strip(),
            'opciones': [opt.strip() for opt in data.get('opciones', '').split(',') if opt.strip()]
        }
        for c in cols:
            if c['nombre'].lower() == new_col['nombre'].lower():
                return jsonify({'error': 'Columna ya existe'}), 400
        cols.append(new_col)
        if config:
            config.valor = safe_json_dumps(cols)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='manual_columns', valor=safe_json_dumps(cols)))
        db.session.commit()
        return jsonify({'success': True})
    if request.method == 'DELETE':
        nombre = request.json.get('nombre')
        cols = [c for c in cols if c['nombre'] != nombre]
        if config:
            config.valor = safe_json_dumps(cols)
            db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/dashboard_charts', methods=['GET', 'POST'])
@login_required
def api_dashboard_charts():
    pid = session.get('current_proyecto_id')
    config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_charts').first()
    
    if request.method == 'GET':
        charts = json.loads(config.valor) if config else []
        return jsonify(charts)
        
    if request.method == 'POST':
        if session.get('rol') in ['demo', 'gestor']:
            return jsonify({'error': 'Rol sin permisos para modificar el dashboard.'}), 403
        charts = request.json # Expecting an array of chart objects
        if config:
            config.valor = safe_json_dumps(charts)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='saved_dashboard_charts', valor=safe_json_dumps(charts)))
        db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/dashboard_kpis', methods=['GET', 'POST'])
@login_required
def api_dashboard_kpis():
    pid = session.get('current_proyecto_id')
    config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_kpis').first()
    
    if request.method == 'GET':
        kpis = json.loads(config.valor) if config else []
        return jsonify(kpis)
        
    if request.method == 'POST':
        if session.get('rol') in ['demo', 'gestor']:
            return jsonify({'error': 'Rol sin permisos para modificar el dashboard.'}), 403
        kpis = request.json
        if config:
            config.valor = safe_json_dumps(kpis)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='saved_dashboard_kpis', valor=safe_json_dumps(kpis)))
        db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/all_columns')
@login_required
def api_master_all_columns():
    pid = session.get('current_proyecto_id')
    schema_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
    manual_cfg = AppConfig.query.filter_by(proyecto_id=pid, clave='manual_columns').first()
    
    cols = set()
    if schema_cfg:
        try: cols.update(json.loads(schema_cfg.valor))
        except: pass
    if manual_cfg:
        try:
            m_list = json.loads(manual_cfg.valor)
            for m in m_list:
                cols.add(m['nombre'])
        except: pass
        
    return jsonify(sorted(list(cols)))

@app.route('/api/master/dashboard_filters', methods=['GET', 'POST'])
@login_required
def api_dashboard_filters():
    pid = session.get('current_proyecto_id')
    config = AppConfig.query.filter_by(proyecto_id=pid, clave='saved_dashboard_filters').first()
    
    if request.method == 'GET':
        filters = json.loads(config.valor) if config else []
        return jsonify(filters)
        
    if request.method == 'POST':
        if session.get('rol') in ['demo', 'gestor']:
            return jsonify({'error': 'Rol sin permisos para modificar filtros del dashboard.'}), 403
        filters = request.json
        if config:
            config.valor = safe_json_dumps(filters)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='saved_dashboard_filters', valor=safe_json_dumps(filters)))
        db.session.commit()
        return jsonify({'success': True})

@app.route('/api/config/consolidation', methods=['GET', 'POST'])
@login_required
def api_config_consolidation():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo' and request.method != 'GET':
        return jsonify({'error': 'Rol DEMO no tiene permisos para modificar la configuración de consolidación.'}), 403
    config = AppConfig.query.filter_by(proyecto_id=pid, clave='consolidation_config').first()
    
    if request.method == 'GET':
        return jsonify(json.loads(config.valor) if config else {})
        
    if request.method == 'POST':
        data = request.json
        if config:
            config.valor = safe_json_dumps(data)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='consolidation_config', valor=safe_json_dumps(data)))
        db.session.commit()
        return jsonify({'success': True})

@app.route('/api/master/template/<tipo>')
@login_required
def api_master_template(tipo):
    output = io.BytesIO()
    if tipo == 'filtros':
        df = pd.DataFrame(columns=['columna', 'valor'])
        filename = "Plantilla_Filtros.xlsx"
    elif tipo == 'tablas':
        df = pd.DataFrame(columns=['columna_criterio', 'valor_criterio', 'nueva_columna', 'nuevo_valor'])
        filename = "Plantilla_Cruces.xlsx"
    else:
        return "Tipo no válido", 400
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Plantilla')
    
    output.seek(0)
    from flask import send_file
    return send_file(output, download_name=filename, as_attachment=True)

@app.route('/api/master/bulk_import/<tipo>', methods=['POST'])
@login_required
def api_master_bulk_import(tipo):
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para realizar importaciones masivas.'}), 403
    if 'file' not in request.files:
        return jsonify({'error': 'No se subió ningún archivo'}), 400
    
    file = request.files['file']
    try:
        df = pd.read_excel(file, dtype=str).fillna('')
        df.columns = [c.strip().lower() for c in df.columns]
        
        added = 0
        if tipo == 'filtros':
            required = ['columna', 'valor']
            if not all(c in df.columns for c in required):
                return jsonify({'error': f'Columnas faltantes. Se requiere: {required}'}), 400
            
            for _, row in df.iterrows():
                try:
                    c, v = row['columna'].strip(), row['valor'].strip()
                    if not c or not v: continue
                    # Check duplicate
                    exists = FiltroMaestro.query.filter_by(proyecto_id=pid, columna=c, valor=v).first()
                    if not exists:
                        nuevo = FiltroMaestro(proyecto_id=pid, columna=c, valor=v)
                        db.session.add(nuevo)
                        added += 1
                except: continue
        
        elif tipo == 'tablas':
            required = ['columna_criterio', 'valor_criterio', 'nueva_columna', 'nuevo_valor']
            if not all(c in df.columns for c in required):
                return jsonify({'error': f'Columnas faltantes. Se requiere: {required}'}), 400
            
            for _, row in df.iterrows():
                try:
                    cc, vc = row['columna_criterio'].strip(), row['valor_criterio'].strip()
                    nc, nv = row['nueva_columna'].strip(), row['nuevo_valor'].strip()
                    if not all([cc, vc, nc, nv]): continue
                    nuevo = TablaMaestra(proyecto_id=pid, columna_criterio=cc, valor_criterio=vc, nueva_columna=nc, nuevo_valor=nv)
                    db.session.add(nuevo)
                    added += 1
                except: continue
        
        db.session.commit()
        return jsonify({'success': True, 'added': added})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/rows/update', methods=['POST'])
@login_required
def api_rows_update():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para actualizar datos.'}), 403
    try:
        data = request.json
        key_val = data.get('key')
        field = data.get('field')
        value = data.get('value')
        if not all([key_val, field]): return jsonify({'error': 'Missing data'}), 400
        record = NucleusData.query.filter_by(proyecto_id=pid, key_value=str(key_val)).first()
        if not record: return jsonify({'error': 'Record not found'}), 404
        row_dict = json.loads(record.data_json)
        
        # Guardar en el historial de cambios (solo si el valor realmente cambió)
        valor_anterior = row_dict.get(field, '')
        if str(valor_anterior) != str(value):
            historial = HistorialCambios(
                proyecto_id=pid,
                usuario_id=session.get('user_id'),
                username=session.get('username'),
                key_value=str(key_val),
                campo_modificado=field,
                valor_anterior=str(valor_anterior),
                valor_nuevo=str(value)
            )
            db.session.add(historial)
            if str(field).strip() == 'Estado de la tarea (WO State)':
                row_dict['FECHA CAMBIO ESTADO'] = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        
        row_dict[field] = value
        row_dict['_ultimo_usuario_manual'] = session.get('username')
        row_dict['_fecha_ultima_act_manual'] = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        
        # --- Instant Logic: Re-apply TablaMaestra rules for this row ---
        tablas = TablaMaestra.query.filter_by(proyecto_id=pid).all()
        for t in tablas:
            t_cols = [c.strip() for c in t.columna_criterio.split(',')]
            t_vals = [v.strip() for v in t.valor_criterio.split(',')]
            
            match = True
            for c, v in zip(t_cols, t_vals):
                if str(row_dict.get(c, '')) != v:
                    match = False
                    break
            if match:
                row_dict[t.nueva_columna] = t.nuevo_valor

        # --- Re-apply Reglas de Estado Manual ---
        reglas_manuales = ReglaEstadoManual.query.filter_by(proyecto_id=pid).all()
        for r in reglas_manuales:
            r_cols = [c.strip() for c in r.columna_criterio.split(',')]
            r_vals = [v.strip() for v in r.valor_criterio.split(',')]
            
            match = True
            for c, v in zip(r_cols, r_vals):
                if str(row_dict.get(c, '')) != v:
                    match = False
                    break
            if match:
                row_dict[r.columna_manual] = r.nuevo_valor

        record.data_json = safe_json_dumps(row_dict)
        db.session.commit()
        
        # Inject KPI calculations for instant feedback
        rows_injected, _ = inject_kpis(pid, [row_dict])
        final_row = rows_injected[0]
        
        # Update schema if new tracking columns
        config_schema = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
        if config_schema:
            schema_cols = set(json.loads(config_schema.valor))
            nuevas_cols = {'_ultimo_usuario_manual', '_fecha_ultima_act_manual'}
            if not nuevas_cols.issubset(schema_cols):
                updated_schema = list(schema_cols.union(nuevas_cols))
                config_schema.valor = safe_json_dumps(updated_schema)
                db.session.commit()
        
        return jsonify({'success': True, 'newData': final_row})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/rows/add', methods=['POST'])
@login_required
def api_rows_add():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para añadir registros.'}), 403
    try:
        data = request.json
        key_val = str(data.get('key', '')).strip()
        
        # Auto-generate key if none provided (proyecto sin clave primaria definida)
        if not key_val:
            existing = NucleusData.query.filter_by(proyecto_id=pid).all()
            nums = []
            for rec in existing:
                try:
                    nums.append(int(float(rec.key_value)))
                except (ValueError, TypeError):
                    pass
            key_val = str(max(nums) + 1) if nums else '1'
        else:
            # Check duplicate
            exists = NucleusData.query.filter_by(proyecto_id=pid, key_value=key_val).first()
            if exists: return jsonify({'error': f'El registro con ID {key_val} ya existe.'}), 400
        
        # Create record (with optional full data)
        row_data = data.get('data') or {}
        if not isinstance(row_data, dict):
            row_data = {}
        new_record = NucleusData(proyecto_id=pid, key_value=key_val, data_json=json.dumps(row_data))
        db.session.add(new_record)
        db.session.commit()
        
        return jsonify({'success': True, 'key': key_val})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/rows/delete', methods=['POST'])
@login_required
def api_rows_delete():
    pid = session.get('current_proyecto_id')
    if session.get('rol') in ['demo', 'gestor']:
        return jsonify({'error': 'No tienes permisos para eliminar registros.'}), 403
    try:
        data = request.json
        keys = data.get('keys', [])
        if not keys: return jsonify({'error': 'No se especificaron registros para eliminar'}), 400
        
        # Delete records
        NucleusData.query.filter(NucleusData.proyecto_id == pid, NucleusData.key_value.in_(keys)).delete(synchronize_session=False)
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# --- WO DETAIL (Modal de detalle del registro) ---
@app.route('/api/wo/meta', methods=['GET'])
@login_required
def api_wo_meta():
    pid = session.get('current_proyecto_id')
    try:
        proy = db.session.get(Proyecto, pid)
        proy_nombre = proy.nombre if proy else ''

        # Opciones de SERVICIO (configurables por proyecto)
        scfg = AppConfig.query.filter_by(proyecto_id=pid, clave='servicio_opciones').first()
        if scfg:
            try:
                servicios = json.loads(scfg.valor)
                if not isinstance(servicios, list):
                    servicios = []
            except Exception:
                servicios = []
        else:
            servicios = ['PREVENTIVO', 'CORRECTIVO', 'PREDICTIVO', 'ABASTECIMIENTO DE COMBUSTIBLE',
                         'ADICIONALES', 'CORTE PROGRAMADO', 'TRABAJO PROGRAMADO']

        # Técnicos activos desde DATAPER, filtrados por el PROYECTO actual
        tecnicos = []
        dataper = Proyecto.query.filter_by(nombre='Dataper').first()
        if dataper:
            seen = set()
            for r in NucleusData.query.filter_by(proyecto_id=dataper.id).all():
                try:
                    d = json.loads(r.data_json)
                except Exception:
                    continue
                est = str(d.get('ESTADO') or '').strip().upper()
                if est and est != 'ACTIVO':
                    continue
                pr = str(d.get('PROYECTO') or '').strip()
                if proy_nombre and pr and pr != proy_nombre:
                    continue
                t = str(d.get('TECNICO') or '').strip()
                if not t or t in seen:
                    continue
                seen.add(t)
                tecnicos.append({'nombre': t, 'contrata': str(d.get('CONTRATA') or '').strip()})
        tecnicos.sort(key=lambda x: x['nombre'])

        # Materiales desde MATERIAL, filtrados por el PROYECTO actual
        materiales = []
        material_proy = Proyecto.query.filter_by(nombre='Material').first()
        if material_proy:
            seen = set()
            for r in NucleusData.query.filter_by(proyecto_id=material_proy.id).all():
                try:
                    d = json.loads(r.data_json)
                except Exception:
                    continue
                pr = str(d.get('PROYECTO') or '').strip()
                if proy_nombre and pr and pr != proy_nombre:
                    continue
                desc = str(d.get('DESCRIPCION_MATERIAL') or '').strip()
                if not desc or desc in seen:
                    continue
                seen.add(desc)
                materiales.append({
                    'codigo': str(d.get('COD_MATERIAL') or '').strip(),
                    'descripcion': desc,
                    'tipo': str(d.get('TIPO') or '').strip(),
                    'um': str(d.get('UM') or '').strip()
                })
        materiales.sort(key=lambda x: x['descripcion'])

        return jsonify({'success': True, 'servicios': servicios, 'tecnicos': tecnicos,
                        'materiales': materiales, 'proyecto': proy_nombre})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/wo/servicios', methods=['POST'])
@login_required
def api_wo_servicios():
    pid = session.get('current_proyecto_id')
    if session.get('rol') not in ['admin', 'supervisor']:
        return jsonify({'error': 'No tienes permisos para editar servicios.'}), 403
    try:
        data = request.json or {}
        opciones = data.get('opciones', [])
        if not isinstance(opciones, list):
            return jsonify({'error': 'Formato inválido'}), 400
        opciones = [str(o).strip() for o in opciones if str(o).strip()]
        scfg = AppConfig.query.filter_by(proyecto_id=pid, clave='servicio_opciones').first()
        if scfg:
            scfg.valor = json.dumps(opciones, ensure_ascii=False)
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='servicio_opciones', valor=json.dumps(opciones, ensure_ascii=False)))
        db.session.commit()
        return jsonify({'success': True, 'servicios': opciones})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/wo/historial', methods=['GET'])
@login_required
def api_wo_historial():
    pid = session.get('current_proyecto_id')
    key_val = (request.args.get('key') or '').strip()
    if not key_val:
        return jsonify({'error': 'Falta el identificador del WO'}), 400
    try:
        rows = (HistorialCambios.query
                .filter_by(proyecto_id=pid, key_value=key_val)
                .order_by(HistorialCambios.fecha.desc(), HistorialCambios.id.desc())
                .all())
        items = [{
            'campo': h.campo_modificado,
            'valor_anterior': h.valor_anterior or '',
            'valor_nuevo': h.valor_nuevo or '',
            'username': h.username,
            # La fecha se guarda en UTC y se convierte a hora de Perú (UTC-5).
            'fecha': (h.fecha - timedelta(hours=5)).isoformat() if h.fecha else None
        } for h in rows]
        return jsonify({'success': True, 'historial': items})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/wo/informe')
@login_required
def api_wo_informe():
    """Genera un informe Word (.docx) con el estado actual e historia de un WO."""
    from flask import make_response
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    pid = session.get('current_proyecto_id')
    key_val = (request.args.get('key') or '').strip()
    if not key_val:
        return jsonify({'error': 'Falta el identificador del WO'}), 400

    record = NucleusData.query.filter_by(proyecto_id=pid, key_value=key_val).first()
    if not record:
        return jsonify({'error': 'Registro no encontrado'}), 404
    row = json.loads(record.data_json)
    row['_key'] = key_val

    proy = db.session.get(Proyecto, pid)
    proy_nombre = proy.nombre if proy else ''

    def get_val(partial, default=''):
        for k, v in row.items():
            if partial.lower() in str(k).lower():
                return v if v is not None else default
        return default

    def peru_fmt(s):
        if not s:
            return ''
        try:
            dt = datetime.strptime(str(s).strip(), '%Y-%m-%d %H:%M:%S')
        except Exception:
            try:
                dt = datetime.fromisoformat(str(s).strip().replace('T', ' ')[:19])
            except Exception:
                return str(s)
        dt = dt - timedelta(hours=5)
        return dt.strftime('%d/%m/%Y %H:%M:%S')

    # Datos clave del WO
    wo_number = str(get_val('Número de WO', row.get('_key', key_val))).strip() or key_val
    estado = str(get_val('wo state', get_val('estado'))).strip()
    tecnico = str(get_val('tecnico asignado')).strip()
    # Solo para el informe: la contrata siempre es COBRA y el técnico lleva (COBRA).
    contrata_info = 'COBRA'
    tecnico_info = f'{tecnico} (COBRA)' if tecnico else ''

    # Materiales (campo MATERIALES JSON)
    materiales = []
    raw_mat = get_val('materiales')
    if isinstance(raw_mat, str):
        try:
            arr = json.loads(raw_mat)
            if isinstance(arr, list):
                materiales = arr
        except Exception:
            materiales = []
    elif isinstance(raw_mat, list):
        materiales = raw_mat

    # Cotización (montos + gastos JSON)
    monto_aprobado = str(get_val('monto aprobado') or '').strip()
    monto_gastando = str(get_val('monto gastando') or '').strip()
    gastos = []
    raw_gastos = get_val('cotizacion_gastos')
    if isinstance(raw_gastos, str):
        try:
            arr = json.loads(raw_gastos)
            if isinstance(arr, list):
                gastos = arr
        except Exception:
            gastos = []
    elif isinstance(raw_gastos, list):
        gastos = raw_gastos

    # Bitácora general
    bitacora = str(get_val('bitácora') or get_val('bitacora')).strip()

    # Evidencia: URLs guardadas + bitácoras por tipo
    ev_tipos = [
        ('INICIO', 'Evidencia de Inicio'),
        ('PROCESO', 'Evidencia de Proceso'),
        ('CIERRE', 'Evidencia de Cierre'),
    ]
    evidencia = []
    for tkey, tlabel in ev_tipos:
        raw_urls = row.get(f'_EVIDENCIA_{tkey}', '')
        urls = []
        if isinstance(raw_urls, str):
            try:
                arr = json.loads(raw_urls)
                if isinstance(arr, list):
                    urls = [u for u in arr if u]
            except Exception:
                urls = []
        elif isinstance(raw_urls, list):
            urls = [u for u in raw_urls if u]
        evidencia.append({
            'tipo': tkey,
            'label': tlabel,
            'urls': urls,
            'bitacora': str(row.get(f'_BITACORA_{tkey}', '') or '').strip()
        })

    # Historial de cambios
    hist = (HistorialCambios.query
            .filter_by(proyecto_id=pid, key_value=key_val)
            .order_by(HistorialCambios.fecha.desc(), HistorialCambios.id.desc())
            .all())

    # --- Construcción del documento ---
    doc = Document()

    # Estilos base
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    # Encabezado: empresa + título
    h = doc.add_paragraph()
    run = h.add_run('COBRA PERÚ')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    srun = sub.add_run('Informe de Avería — Detalle e Historial')
    srun.font.size = Pt(13)
    srun.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Línea separadora
    sep = doc.add_paragraph()
    srun2 = sep.add_run('—' * 40)
    srun2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # Datos generales
    doc.add_heading('1. Datos Generales', level=1)
    datos = [
        ('N° de WO', wo_number),
        ('Estado', estado),
        ('Proyecto', proy_nombre),
        ('Técnico asignado', tecnico_info),
        ('Contrata', contrata_info),
        ('Servicio', str(get_val('servicio')).strip()),
        ('Ciudad', str(get_val('ciudad')).strip()),
        ('Motivo de avería', str(get_val('motivo de avería')).strip()),
        ('Solución', str(get_val('solución')).strip()),
        ('¿Se instaló mufas?', str(get_val('se instaló mufas')).strip()),
        ('Inicio de parada', peru_fmt(str(get_val('inicio de parada')).strip())),
        ('Fin de parada', peru_fmt(str(get_val('fin de parada')).strip())),
        ('Monto aprobado', str(get_val('monto aprobado')).strip()),
        ('Monto gastando', str(get_val('monto gastando')).strip()),
        ('¿Requiere correctivo final?', str(get_val('requiere correctivo final')).strip()),
        ('Fecha de creación (WO)', peru_fmt(str(get_val('Fecha de creación (WO Creation date)')).strip())),
        ('Fecha closed', peru_fmt(str(get_val('Fecha closed')).strip())),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in datos:
        cells = table.add_row().cells
        cells[0].width = Inches(2.3)
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].paragraphs[0].add_run(v or '—')

    # Estado actual de las columnas del registro (todo lo demás con valor)
    resto = {}
    for k, v in row.items():
        if k.startswith('_'):
            continue
        if k in ('MATERIALES', 'MATERIAL USADO'):
            continue
        s = '' if v is None else str(v)
        if s.strip():
            resto[k] = s.strip()
    if resto:
        doc.add_heading('2. Estado Actual del Registro', level=1)
        t2 = doc.add_table(rows=0, cols=2)
        t2.style = 'Light Grid Accent 1'
        for k, v in resto.items():
            cells = t2.add_row().cells
            cells[0].width = Inches(2.3)
            cells[0].paragraphs[0].add_run(k).bold = True
            cells[1].paragraphs[0].add_run(v)

    # Cotización
    if monto_aprobado or monto_gastando or gastos:
        doc.add_heading('3. Cotización', level=1)
        if monto_aprobado:
            p = doc.add_paragraph()
            r = p.add_run('Monto aprobado: ')
            r.bold = True
            p.add_run(f'S/ {monto_aprobado}')
        if gastos:
            tg = doc.add_table(rows=1, cols=2)
            tg.style = 'Light Grid Accent 1'
            for i, t in enumerate(['Concepto', 'Monto (S/)']):
                tg.rows[0].cells[i].paragraphs[0].add_run(t).bold = True
            for g in gastos:
                cells = tg.add_row().cells
                cells[0].text = str(g.get('descripcion', '') or '')
                cells[1].text = str(g.get('monto', '') or '')
            if monto_gastando:
                p = doc.add_paragraph()
                r = p.add_run('Monto gastado total: ')
                r.bold = True
                p.add_run(f'S/ {monto_gastando}')

    # Materiales usados
    if materiales:
        doc.add_heading('4. Materiales Usados', level=1)
        tm = doc.add_table(rows=1, cols=5)
        tm.style = 'Light Grid Accent 1'
        hdr = tm.rows[0].cells
        for i, t in enumerate(['Código', 'Descripción', 'Tipo', 'UM', 'Cantidad']):
            hdr[i].paragraphs[0].add_run(t).bold = True
        for m in materiales:
            cells = tm.add_row().cells
            cells[0].text = str(m.get('codigo', '') or '')
            cells[1].text = str(m.get('descripcion', '') or '')
            cells[2].text = str(m.get('tipo', '') or '')
            cells[3].text = str(m.get('um', '') or '')
            cells[4].text = str(m.get('cantidad', '') or '')

    # Bitácora
    if bitacora:
        doc.add_heading('5. Bitácora', level=1)
        doc.add_paragraph(bitacora)

    # Historial de cambios
    if hist:
        doc.add_heading('6. Historial de Cambios', level=1)
        th = doc.add_table(rows=1, cols=4)
        th.style = 'Light Grid Accent 1'
        for i, t in enumerate(['Fecha', 'Usuario', 'Campo', 'Cambio (anterior → nuevo)']):
            th.rows[0].cells[i].paragraphs[0].add_run(t).bold = True
        for hc in hist:
            cells = th.add_row().cells
            cells[0].text = peru_fmt((hc.fecha - timedelta(hours=0)).strftime('%Y-%m-%d %H:%M:%S')) if hc.fecha else ''
            cells[1].text = hc.username
            cells[2].text = hc.campo_modificado
            cambio = f'{hc.valor_anterior or ""} → {hc.valor_nuevo or ""}'
            cells[3].text = cambio

    # Evidencia fotográfica
    doc.add_heading('7. Evidencia Fotográfica', level=1)
    for ev in evidencia:
        if not ev['urls'] and not ev['bitacora']:
            continue
        doc.add_heading(ev['label'], level=2)
        if ev['bitacora']:
            p = doc.add_paragraph()
            r = p.add_run('Nota: ')
            r.bold = True
            p.add_run(ev['bitacora'])
        for url in ev['urls']:
            nombre = os.path.basename(str(url).split('?')[0])
            img_bytes = None
            if evidencia_usa_b2():
                try:
                    obj = b2_cliente().get_object(Bucket=app.config['B2_BUCKET'], Key=f'{key_val}/{nombre}')
                    img_bytes = obj['Body'].read()
                except Exception:
                    img_bytes = None
            else:
                ruta = os.path.join(evidencia_folder(pid, key_val), nombre)
                if os.path.exists(ruta):
                    with open(ruta, 'rb') as f:
                        img_bytes = f.read()
            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
                except Exception:
                    doc.add_paragraph(f'(Imagen no embebible: {nombre})')
            else:
                doc.add_paragraph(f'(Imagen no encontrada: {nombre})')

    # Pie de documento
    doc.add_paragraph()
    pie = doc.add_paragraph()
    prun = pie.add_run('Generado por Nucleus — COBRA PERÚ')
    prun.font.size = Pt(8)
    prun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Salida
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    safe_key = secure_filename(key_val) or 'WO'
    response = make_response(output.read())
    response.headers['Content-Disposition'] = f'attachment; filename=Informe_{safe_key}.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response

@app.route('/api/rows/bulk_update', methods=['POST'])
@login_required
def api_rows_bulk_update():
    pid = session.get('current_proyecto_id')
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para actualizar datos.'}), 403
    try:
        data = request.json
        key_val = data.get('key')
        updates = data.get('data')
        if not key_val or not isinstance(updates, dict):
            return jsonify({'error': 'Datos incompletos'}), 400
        record = NucleusData.query.filter_by(proyecto_id=pid, key_value=str(key_val)).first()
        if not record: return jsonify({'error': 'Registro no encontrado'}), 404
        row_dict = json.loads(record.data_json)
        old_state = str(row_dict.get('Estado de la tarea (WO State)', '')).strip()

        for field, value in updates.items():
            valor_anterior = row_dict.get(field, '')
            if str(valor_anterior) != str(value):
                historial = HistorialCambios(
                    proyecto_id=pid,
                    usuario_id=session.get('user_id'),
                    username=session.get('username'),
                    key_value=str(key_val),
                    campo_modificado=field,
                    valor_anterior=str(valor_anterior),
                    valor_nuevo=str(value)
                )
                db.session.add(historial)
            row_dict[field] = value
        row_dict['_ultimo_usuario_manual'] = session.get('username')
        row_dict['_fecha_ultima_act_manual'] = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')

        new_state = str(row_dict.get('Estado de la tarea (WO State)', '')).strip()
        if new_state and new_state != old_state:
            row_dict['FECHA CAMBIO ESTADO'] = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')

        # --- Instant Logic: TablaMaestra ---
        tablas = TablaMaestra.query.filter_by(proyecto_id=pid).all()
        for t in tablas:
            t_cols = [c.strip() for c in t.columna_criterio.split(',')]
            t_vals = [v.strip() for v in t.valor_criterio.split(',')]
            match = True
            for c, v in zip(t_cols, t_vals):
                if str(row_dict.get(c, '')) != v:
                    match = False
                    break
            if match:
                row_dict[t.nueva_columna] = t.nuevo_valor

        # --- Instant Logic: Reglas de Estado Manual ---
        reglas_manuales = ReglaEstadoManual.query.filter_by(proyecto_id=pid).all()
        for r in reglas_manuales:
            r_cols = [c.strip() for c in r.columna_criterio.split(',')]
            r_vals = [v.strip() for v in r.valor_criterio.split(',')]
            match = True
            for c, v in zip(r_cols, r_vals):
                if str(row_dict.get(c, '')) != v:
                    match = False
                    break
            if match:
                row_dict[r.columna_manual] = r.nuevo_valor

        record.data_json = safe_json_dumps(row_dict)
        db.session.commit()

        rows_injected, _ = inject_kpis(pid, [row_dict])
        final_row = rows_injected[0]

        # Update schema if new tracking columns
        config_schema = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
        if config_schema:
            schema_cols = set(json.loads(config_schema.valor))
            nuevas_cols = {'_ultimo_usuario_manual', '_fecha_ultima_act_manual', 'FECHA CAMBIO ESTADO'}
            if not nuevas_cols.issubset(schema_cols):
                updated_schema = list(schema_cols.union(nuevas_cols))
                config_schema.valor = safe_json_dumps(updated_schema)
                db.session.commit()

        return jsonify({'success': True, 'newData': final_row})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/config/init_manual', methods=['POST'])
@login_required
def api_config_init_manual():
    pid = session.get('current_proyecto_id')
    if session.get('rol') not in ['admin', 'supervisor']:
        return jsonify({'error': 'No tienes permisos para inicializar proyectos.'}), 403
    try:
        data = request.json
        pk_name = str(data.get('primary_key', '')).strip()
        if not pk_name: return jsonify({'error': 'El nombre de la columna principal es requerido.'}), 400
        
        # Initialize Primary Key and empty Schema
        config_pk = AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').first()
        if config_pk:
            config_pk.valor = pk_name
        else:
            db.session.add(AppConfig(proyecto_id=pid, clave='primary_key', valor=pk_name))
            
        config_schema = AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').first()
        if not config_schema:
            db.session.add(AppConfig(proyecto_id=pid, clave='app_schema', valor=json.dumps([])))
            
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/clean', methods=['POST'])
@login_required
def api_clean():
    if session.get('rol') in ['gestor', 'demo']:
        return jsonify({'error': 'No tienes permisos para esta acción.'}), 403
        
    pid = session.get('current_proyecto_id')
    try:
        NucleusData.query.filter_by(proyecto_id=pid).delete()
        AppConfig.query.filter_by(proyecto_id=pid, clave='app_schema').delete()
        AppConfig.query.filter_by(proyecto_id=pid, clave='primary_key').delete()
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# --- Evidencia fotográfica (Tiempo Inicio / Proceso / Cierre) ---
EVIDENCIA_TIPOS = ('inicio', 'proceso', 'cierre')
EVIDENCIA_MAX_POR_TIPO = 5
EVIDENCIA_EXT_ALLOWED = {'.jpg', '.jpeg', '.png', '.webp', '.gif', '.bmp', '.heic'}

def evidencia_folder(pid, key):
    return os.path.join(app.config['EVIDENCIA_DIR'], str(pid), secure_filename(str(key)))

def evidencia_limpiar_slot(pid, key, tipo, indice):
    folder = evidencia_folder(pid, key)
    for old in glob.glob(os.path.join(folder, f'{tipo}_{int(indice)}.*')):
        try:
            os.remove(old)
        except Exception:
            pass

def evidencia_comprimir(ruta, calidad=None, max_lado=None):
    """Comprime y redimensiona una imagen para ahorrar almacenamiento.
    Devuelve True si se comprimió correctamente; si no puede (ej. HEIC),
    deja el archivo original tal cual."""
    calidad = calidad if calidad is not None else app.config['EVIDENCIA_CALIDAD']
    max_lado = max_lado if max_lado is not None else app.config['EVIDENCIA_MAX_LADO']
    try:
        img = Image.open(ruta)
        img = ImageOps.exif_transpose(img)
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGBA')
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1])
            img = bg
        else:
            img = img.convert('RGB')
        w, h = img.size
        lado_max = max(w, h)
        if lado_max > max_lado:
            ratio = max_lado / lado_max
            img = img.resize((max(1, int(w * ratio)), max(1, int(h * ratio))), Image.LANCZOS)
        img.save(ruta, 'JPEG', quality=calidad, optimize=True)
        return True
    except Exception:
        return False

def evidencia_usa_b2():
    return bool(app.config.get('B2_BUCKET') and app.config.get('B2_KEY_ID') and app.config.get('B2_APP_KEY'))

def b2_cliente():
    import boto3
    endpoint = app.config['B2_ENDPOINT_URL'] or f"https://s3.{app.config['B2_REGION']}.backblazeb2.com"
    return boto3.client(
        's3',
        endpoint_url=endpoint,
        aws_access_key_id=app.config['B2_KEY_ID'],
        aws_secret_access_key=app.config['B2_APP_KEY'],
        region_name=app.config['B2_REGION'],
    )

def evidencia_eliminar_b2(key, tipo, indice):
    client = b2_cliente()
    bucket = app.config['B2_BUCKET']
    prefix = f'{key}/{tipo}_{int(indice)}.'
    try:
        resp = client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        for o in resp.get('Contents', []):
            client.delete_object(Bucket=bucket, Key=o['Key'])
    except Exception:
        pass

@app.route('/api/evidencia/subir', methods=['POST'])
@login_required
def api_evidencia_subir():
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos para subir evidencia.'}), 403
    pid = session.get('current_proyecto_id')
    key = (request.form.get('key') or '').strip()
    tipo = (request.form.get('tipo') or '').strip().lower()
    try:
        indice = int(request.form.get('indice'))
    except (TypeError, ValueError):
        return jsonify({'error': 'Índice inválido'}), 400
    file = request.files.get('foto')
    if not key or tipo not in EVIDENCIA_TIPOS:
        return jsonify({'error': 'Datos incompletos'}), 400
    if indice < 0 or indice >= EVIDENCIA_MAX_POR_TIPO:
        return jsonify({'error': 'Índice fuera de rango'}), 400
    if file is None or not file.filename:
        return jsonify({'error': 'No se recibió ningún archivo'}), 400
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in EVIDENCIA_EXT_ALLOWED:
        return jsonify({'error': f'Formato no permitido: {ext}. Usa {", ".join(sorted(EVIDENCIA_EXT_ALLOWED))}'}), 400

    # Guardar a un archivo temporal, comprimir y luego mover/subir.
    fd, ruta_tmp = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        file.save(ruta_tmp)
        nombre = f'{tipo}_{indice}{ext}'
        if evidencia_comprimir(ruta_tmp):
            if not nombre.lower().endswith('.jpg'):
                nueva = ruta_tmp + '.jpg'
                os.rename(ruta_tmp, nueva)
                ruta_tmp = nueva
                nombre = f'{tipo}_{indice}.jpg'

        if evidencia_usa_b2():
            evidencia_eliminar_b2(key, tipo, indice)
            b2_cliente().upload_file(ruta_tmp, app.config['B2_BUCKET'], f'{key}/{nombre}')
        else:
            folder = evidencia_folder(pid, key)
            os.makedirs(folder, exist_ok=True)
            evidencia_limpiar_slot(pid, key, tipo, indice)
            os.replace(ruta_tmp, os.path.join(folder, nombre))

        url = f'/api/evidencia/foto/{pid}/{secure_filename(str(key))}/{nombre}?v={int(time.time())}'
        return jsonify({'success': True, 'url': url})
    finally:
        if os.path.exists(ruta_tmp):
            try:
                os.remove(ruta_tmp)
            except Exception:
                pass

@app.route('/api/evidencia/eliminar', methods=['POST'])
@login_required
def api_evidencia_eliminar():
    if session.get('rol') == 'demo':
        return jsonify({'error': 'Rol DEMO no tiene permisos.'}), 403
    pid = session.get('current_proyecto_id')
    data = request.json or {}
    key = (data.get('key') or '').strip()
    tipo = (data.get('tipo') or '').strip().lower()
    try:
        indice = int(data.get('indice'))
    except (TypeError, ValueError):
        return jsonify({'error': 'Índice inválido'}), 400
    if not key or tipo not in EVIDENCIA_TIPOS or indice < 0 or indice >= EVIDENCIA_MAX_POR_TIPO:
        return jsonify({'error': 'Datos incompletos'}), 400
    if evidencia_usa_b2():
        evidencia_eliminar_b2(key, tipo, indice)
    else:
        evidencia_limpiar_slot(pid, key, tipo, indice)
    return jsonify({'success': True})

@app.route('/api/evidencia/foto/<int:pid>/<path:key>/<path:nombre>')
@login_required
def api_evidencia_foto(pid, key, nombre):
    if pid != session.get('current_proyecto_id'):
        return jsonify({'error': 'Acceso denegado'}), 403
    nombre = os.path.basename(nombre)
    if evidencia_usa_b2():
        try:
            obj = b2_cliente().get_object(Bucket=app.config['B2_BUCKET'], Key=f'{key}/{nombre}')
            data = obj['Body'].read()
            mt = mimetypes.guess_type(nombre)[0] or 'application/octet-stream'
            resp = Response(data, mimetype=mt)
            resp.headers['Cache-Control'] = 'private, max-age=604800, immutable'
            return resp
        except Exception:
            return jsonify({'error': 'No encontrado'}), 404
    folder = evidencia_folder(pid, key)
    return send_from_directory(folder, nombre, max_age=604800)

@app.route('/healthz')
def healthz():
    return 'OK', 200

if __name__ == '__main__':
    app.run(debug=True, port=int(os.environ.get('PORT', 5000)))
